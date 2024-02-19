import customtkinter as ct
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import numpy as np
import pandas as pd
import os
import tkinter as tk 
from tkinter import filedialog
from PIL import Image
from tkinter import ttk
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from itertools import permutations
from docx import Document
from tkinter import Tk
from io import BytesIO
from docx.shared import Inches




color_names = ['red', 'green', 'blue', 'navy' , 'orange', 'purple', 'pink', 'brown', 'cyan', 'magenta',
               'teal', 'lime', 'lavender', 'indigo', 'maroon', 'gold', 'olive', 'yellow']
color_names_prep = ['black','red','black', 'green','black', 'blue','black', 'yellow','black', 'orange','black', 'purple','black', 'pink','black', 'brown','black', 'cyan','black', 'magenta','black',
                    'teal','black', 'lime','black', 'lavender','black', 'indigo','black', 'maroon','black', 'gold','black', 'olive','black', 'navy','black', 'silver']


























# Fonction pour transformer une chaîne de texte en une matrice
def ttm(txt):
    # Divise le texte en lignes
    L = txt.split("\n")
    l, m = [], []

    # Parcourt chaque ligne
    for i in range(len(L)):
        # Ignore les lignes vides ou avec des espaces multiples
        if L[i] == "" or L[i] == " " or L[i] == "  " or L[i] == "   " or L[i] == "    ":
            pass
        else:
            # Ajoute les éléments de la ligne à la liste l
            l.append(L[i].split(","))

    # Convertit les éléments en nombres flottants et construit la matrice
    for e, i in zip(l, range(len(l))):
        m.append([])
        for f, j in zip(e, range(len(e))):
            try:
                m[i].append(float(f))
            except:
                print("erreur de syntaxe")
                return "erreur de syntaxe"

    # Tente de créer un tableau NumPy à partir de la matrice
    try:
        test = np.array(m)
    except:
        return "erreur dans la matrice"

    return m

# Fonction pour tester la validité d'une séquence
def test_seq(SQ, n):
    seq = SQ.copy()

    # Vérifie si la longueur de la séquence est correcte
    if len(seq) != n:
        return "le nombre des jobs n'est pas exacte"

    # Trie la séquence et la compare avec une séquence ordonnée
    seq.sort()
    lst = list(range(n))
    if lst != seq:
        return 'la sequence est erronée'

    return SQ

# Fonction pour convertir une chaîne en liste
def char_to_lst(char):
    if char == "":
        print("le champ des temps de préparation est vide")
        return "le champ des temps de préparation est vide"

    lst1 = char.split("-")
    lst2 = []

    # Divise la chaîne en listes
    for e in lst1:
        lst2.append(e.split("\n"))

    lst3 = []
    lst22 = lst2.copy()

    # Construit une liste de listes de nombres entiers
    for i in range(len(lst22)):
        lst3.append([])
        for j in range(len(lst22[0])):
            if lst2[i][j] == "" or lst2[i][j] == " " or lst2[i][j] == " " or lst2[i][j] == " ":
                pass
            else:
                lst3[i].append(lst2[i][j].split(","))

    lst4 = []

    # Convertit les éléments en nombres entiers
    for i in range(len(lst3)):
        lst4.append([])
        for j in range(len(lst3[i])):
            lst4[i].append([])
            for k in range(len(lst3[i][j])):
                try:
                    lst4[i][j].append(float(lst3[i][j][k]))
                except:
                    print("erreur de syntaxe dans champ des temps de préparation")
                    return "erreur de syntaxe dans champ des temps de préparation"

    try:
        for i in range(len(lst4)):
            np.array(lst4[i])
    except:
        print("erreur de syntaxe dans champ des temps de préparation")
        return "erreur de syntaxe dans champ des temps de préparation"

    return lst4

# Fonction pour générer toutes les combinaisons possibles
def generate_combinations(n):
    numbers = list(range(n))
    combinations = list(permutations(numbers))
    return combinations

# Fonction pour obtenir le chemin d'un fichier via une boîte de dialogue
def get_file_path():
    file_path1 = filedialog.askopenfilename(title="Sélectionner un fichier", filetypes=(("Fichiers Python", "*.txt"), ("python", "*.py")))
    return file_path1

# Fonction pour combiner deux listes
def combine_lists(A, B):
    combined_list = []
    for a in A:
        for b in B:
            combined_list.append(a + b)
    return combined_list

# Fonction pour trouver les combinaisons valides de deux listes
def find_combinations(A, B):
    n = len(A)
    valid_combinations = []

    def generate_combinations(current_combination, a_index):
        if a_index == n:
            valid_combinations.append(current_combination[:])
            return

        value = A[a_index]
        for b_index in range(n):
            if B[b_index] == B[a_index] and b_index not in current_combination:
                current_combination[a_index] = b_index
                generate_combinations(current_combination, a_index + 1)
                current_combination[a_index] = -1

    # Initialise la liste de combinaisons
    current_combination = [-1] * n
    generate_combinations(current_combination, 0)
    R = []
    for combination in valid_combinations:
        result = [A[i] for i in combination]
        R.append(result)
    return R

# Fonction pour transposer une matrice
def transpose_matrix(matrix):
    transposed_matrix = [[0] * len(matrix) for _ in range(len(matrix[0]))]
    for i in range(len(matrix)):
        for j in range(len(matrix[0])):
            transposed_matrix[j][i] = matrix[i][j]
    return transposed_matrix

#fonction pour tester le syntaxt dans l'entrée des temps de delais
def dj_test(dj):
    if dj == '':
        return "chaps des delais vide"
    L  = dj.split(",")
    DJ = []
    try: 
        for i in range(len(L)):
            DJ.append(float(L[i]))
    except:
        return "erreur dans les temps de delai"
    return DJ










class Report:
    def __init__(self, regular=0, prep=0, block=0, prep_block=0, noidel=0, nowait=0, TT=0, Seq=[], dj=[], Matrix=[[]], TP=[[[]]]):
        self.regular = regular
        self.prep = prep
        self.block = block
        self.prep_block = prep_block 
        self.noidel = noidel
        self.nowait = nowait
        self.TT = TT
        self.seq = Seq
        self.dj = dj
        if self.dj==[]:
            self.use_TT =0
        else:
            self.use_TT =1
        self.matrix = Matrix
        self.TP = TP
        self.Title1 = 1
        self.Title2 = 1
        self.Table_num = 1
        self.fig_num = 1
        self.m = len(self.matrix)
        self.n = len(self.matrix[0])
        self.doc = Document()
        self.data_add()
        if self.regular == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+'. PFSP (Permutation Flow Shop Scheduling Problem)').style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.regular_add()
            self.regular_add_without_seq()
        if self.prep == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+". PFSP-SDST (Permutation Flow Shop Scheduling Problem with Sequence Depending Setup Time)").style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.prep_add()
            self.seq_choose(1,0,0,0)
            self.prep_add()
        if self.block == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+". BFSP (Blocking Flow Shop Scheduling Problem)").style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.block_add()
            self.seq_choose(0,1,0,0)
            self.block_add()
        if self.prep_block == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+". BFSP-SDST (Blocking Flow Shop Scheduling Problem with Sequence Depending Setup Time)").style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.prep_block_add()
            self.seq_choose(1,1,0,0)
            self.prep_block_add()
        if self.noidel == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+". NIPFSP \n(No-Idel Permutation Flow Shop Scheduling Problem)").style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.noidel_add()
            self.seq_choose(0,0,0,1)
            self.noidel_add()
        if self.nowait == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+". NWPFSP \n(No-Wait Permutation Flow Shop Scheduling Problem)").style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.nowait_add()
            self.seq_choose(0,0,1,0)
            self.nowait_add()
        self.file_path = get_save_path()
        self.save_doc()

    def data_add(self):
        self.doc.add_heading(str(self.Title1)+'. Data de Problème:', level=1).style = 'Heading1'
        self.Title1+=1
        self.doc.add_heading("    "+str(self.Title2)+'. Tableau de temps de processing').style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("P[i][S[j]]: matrice des temps de processing")
        self.table_P = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_P.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Temps de processing').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_P.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(self.matrix[i-1][j-1])
                elif i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "P"
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(j-1)
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
        if self.TP!=[[[]]]:
            self.doc.add_heading("    "+str(self.Title2)+'. Tableau de temps de préparation').style = 'Heading2'
            self.Title2+=1
            self.doc.add_paragraph("TP[i][S[j1]][S[j2]]: Tenseur des temps de préparation")
            self.table_TP=[]
            for i in range(self.m):
                self.table_TP.append(self.doc.add_table(rows=self.n+1, cols=self.n+1)) 
                self.table_TP[i].style = 'Table Grid'
                caption = self.doc.add_paragraph()
                caption.add_run('Table ' + str(self.Table_num) +  'Temps de préparation sur la machine '+str(i)).bold = True
                caption.style = 'Caption'
                self.Table_num+=1
                for j1 in range(self.n+1):
                    for j2 in range(self.n+1):
                        cell = self.table_TP[i].cell(j1, j2)
                        # Ensure there is at least one paragraph and run in the cell
                        if not cell.paragraphs:
                            cell.add_paragraph()
                        if not cell.paragraphs[0].runs:
                            cell.paragraphs[0].add_run()
                        
                        if j1 != 0 and j2 != 0:
                            cell.paragraphs[0].runs[0].text = str(self.TP[i][j1-1][j2-1])
                        elif j1 == 0 and j2 == 0:
                            cell.paragraphs[0].runs[0].text = "TP"+str(i)
                        elif j1 == 0 and j2 != 0:
                            cell.paragraphs[0].runs[0].text = 'J'+str(j2-1)
                        elif j1 != 0 and j2 == 0:
                            cell.paragraphs[0].runs[0].text = 'J'+str(j1-1)
        if self.dj != []:
            self.doc.add_heading("    "+str(self.Title2)+'. Liste des délais de chaque Job').style = 'Heading2'
            self.Title2+=1
            self.doc.add_paragraph("d[S[j]]: List des délais de chaque job")
            self.table_dj = self.doc.add_table(rows=2, cols=self.n+1)
            self.table_dj.style = 'Table Grid'
            caption = self.doc.add_paragraph()
            caption.add_run('Table ' + str(self.Table_num) + ': Délais de chaque job').bold = True
            caption.style = 'Caption'
            self.Table_num+=1
            for i in range(2):
                for j in range(self.n+1):
                    cell = self.table_dj.cell(i, j)
                    # Ensure there is at least one paragraph and run in the cell
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    if not cell.paragraphs[0].runs:
                        cell.paragraphs[0].add_run()
                    
                    if i == 1 and j != 0:
                        cell.paragraphs[0].runs[0].text = str(self.dj[j-1])
                    elif i == 0 and j != 0:
                        cell.paragraphs[0].runs[0].text = 'J'+str(j-1)
                    elif i == 0 and j == 0:
                        cell.paragraphs[0].runs[0].text = ''
                    elif i == 1 and j == 0:
                        cell.paragraphs[0].runs[0].text = 'd'
        if self.seq != []:
            self.doc.add_heading("    "+str(self.Title2)+'. Séquence spécifique des jobs').style = 'Heading2'
            self.Title2+=1
            self.doc.add_paragraph("S[j]: Séquence spécifique des jobs")
            self.table_seq = self.doc.add_table(rows=1, cols=self.n+1)
            self.table_seq.style = 'Table Grid'
            caption = self.doc.add_paragraph()
            caption.add_run('Table ' + str(self.Table_num) + ': Séquence spécifique des jobs').bold = True
            caption.style = 'Caption'
            self.Table_num+=1
            for i in range(1):
                for j in range(self.n+1):
                    cell = self.table_seq.cell(i, j)
                    # Ensure there is at least one paragraph and run in the cell
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    if not cell.paragraphs[0].runs:
                        cell.paragraphs[0].add_run()
                    
                    if j != 0:
                        cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                    elif j == 0:
                        cell.paragraphs[0].runs[0].text = 'S'
    
    def regular_add_without_seq(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,M=self.matrix,Mprep=self.TP)
        if self.m<=2:
            self.doc.add_heading("    "+str(self.Title2)+". Séquence par algorithme de Jonhson").style = 'Heading2'
            self.Title2+=1
            Results = Problem.Johnson()
            self.seq = Results[-1]
        else:
            self.doc.add_heading("    "+str(self.Title2)+". Séquence par algorithme de CDS").style = 'Heading2'
            self.Title2+=1
            combinations = []
            Cmaxs = []
            Matrixs = []
            for k in range(1,self.m):
                print("k=",k)
                Results = Problem.CDS(k)
                print(Results,'mmmmmmmmmmmmmmmmmmmm')
                Problem = ordonnancement(TT=self.use_TT,dj=self.dj,S=Results[0][-1],do_report=1,M=self.matrix,Mprep=self.TP)
                combination = []
                Cmax=[]
                for seq in Results[0][-1]:
                    print("seq=",seq)
                    Cmax.append(Problem.Cmax(seq))
                    combination.append(seq)
                ind = Cmax.index(min(Cmax))
                Cmaxs.append(Cmax[ind])
                combinations.append(combination[ind])
                Matrixs.append(Results[-1])
            index = Cmaxs.index(min(Cmaxs))
            self.seq = combinations[index]
            print(self.seq,"dsfdsfds")
            matrix = Matrixs[index]
            k = index+1
            Results = Problem.CDS(k)
            Results = [Results[0][0],Results[0][1],Results[0][2],Results[0][3],Results[0][4]]
            self.doc.add_paragraph("après le teste de tous les k possibles de k = 1,...,"+str(self.m-1)+"on a trouvé que le meilleur k qui nous donne Cmax le plus optimale est k="+str(k))
            self.doc.add_paragraph("P[i][S[j]]: matrice des temps de processing trouvée par CDS")
            self.table_P_CDS = self.doc.add_table(rows=3, cols=self.n+1)
            self.table_P_CDS.style = 'Table Grid'
            caption = self.doc.add_paragraph()
            caption.add_run('Table ' + str(self.Table_num) + ': Temps de processing des deux machines fictives pou k optimale').bold = True
            caption.style = 'Caption'
            self.Table_num+=1
            for i in range(3):
                for j in range(self.n+1):
                    cell = self.table_P_CDS.cell(i, j)
                    # Ensure there is at least one paragraph and run in the cell
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    if not cell.paragraphs[0].runs:
                        cell.paragraphs[0].add_run()
                    
                    if i != 0 and j != 0:
                        cell.paragraphs[0].runs[0].text = str(matrix[i-1][j-1])
                    elif i == 0 and j == 0:
                        cell.paragraphs[0].runs[0].text = "P_CDS"
                    elif i == 0 and j != 0:
                        cell.paragraphs[0].runs[0].text = 'J'+str(j-1)
                    elif i != 0 and j == 0:
                        cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
        
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,M=self.matrix,Mprep=self.TP)
        self.doc.add_paragraph("S[j]: Séquence trouvée par l'algotrithme de Jonhson")
        self.doc.add_paragraph("U[j]: Jobs tel que P[0][S[j]]<P[1][S[j]]:"+str(Results[0]))
        self.doc.add_paragraph("V[j]: Jobs tel que P[0][S[j]]>=P[1][S[j]]:"+str(Results[1]))
        self.doc.add_paragraph("U_SPT[j]: U[j] avec Jobs ordonancer par SPT(short processing Time), temps de processing croissant:"+str(Results[2]))
        self.doc.add_paragraph("V_LPT[j]: V[j] avec Jobs ordonancer par LPT(long processing Time), temps de processing décroissant:"+str(Results[3]))
        self.doc.add_paragraph("Alors la séquence trouvée par Jonhson est la suivante ")
        self.table_seq_jonhson = self.doc.add_table(rows=1, cols=self.n+1)
        self.table_seq_jonhson.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Séquence trouvée par l'algotrithme de Jonhson").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(1):
            for j in range(self.n+1):
                cell = self.table_seq_jonhson.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif j == 0:
                    cell.paragraphs[0].runs[0].text = 'S'
        self.regular_add()

    def seq_choose(self,p,b,now,noi):
        print(p,"ppppppppppppp")
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,M=self.matrix,Mprep=self.TP,prep=p,block=b,noidel=noi,nowait=now)
        is_TT = 0
        if self.dj!=[]:
            choosed = Problem.choosed_combinations_by_TT_Cmax()
            is_TT=1
            if type(choosed)==str:
                choosed = Problem.choosed_combinations()
                is_TT=0
        else:
            choosed = Problem.choosed_combinations()
        print(choosed,'cccccccccccccccccccccc')
        self.seq = choosed[1][0]
        if is_TT==0:
            self.doc.add_heading("    "+str(self.Title2)+". Séquence par Test de tous les permutations possibles avec min(Cmax)").style = 'Heading2'
        else:
            self.doc.add_heading("    "+str(self.Title2)+". Séquence par Test de tous les permutations possibles avec min(Cmax) et min(TT)").style = 'Heading2'
        self.Title2+=1
        self.table_seq_comb = self.doc.add_table(rows=1, cols=self.n+1)
        self.table_seq_comb.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Séquence trouvée par l'algotrithme de Test de toutes les séquences").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(1):
            for j in range(self.n+1):
                cell = self.table_seq_comb.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif j == 0:
                    cell.paragraphs[0].runs[0].text = 'S'

    def regular_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de PFSP (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("C[i][S[j]]: matrice des dates de fin de chaque job dans chaque machine \n")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrice des dates des fins et des débuts de chaque job sur chaque machine").style = 'Heading2'
        self.Title2+=1
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine PFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])
                                
        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine PFSP \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine PFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines PFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[0][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[1][i-1]*100,2))+"%"
        Figure = plot_Performances(Performances)
        self.doc.add_picture(Figure, width=Inches(5)) 

        
        self.doc.add_heading("    "+str(self.Title2)+". Temps d'attedre de chaque job entre chaque deux machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TW: (waiting Time) temps d'attedre de chaque job entre chaque deux machines")
        waitting_times = Problem.TW()
        self.table_TW = self.doc.add_table(rows=self.m, cols=self.n+1)
        self.table_TW.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Temps d'attendre de chaque job entre chaque deux machines PFSP").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m):
            for j in range(self.n+1):
                cell = self.table_TW.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "TW" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)+'/'+str(i)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(waitting_times[0][i-1][j-1])
        
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),0,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))
    
    def prep_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP,prep=1)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de PFSP-SDST (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("C[i][S[j]]: matrice des dates de fin de chaque job dans chaque machine \nM[i][S[j]][S[j]]: Tenseur des temps de préparation\n ")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrice des dates des fins et des débuts de chaque job sur chaque machine").style = 'Heading2'
        self.Title2+=1
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine PFSP-SDST').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])
                                
        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine PFSP-SDST \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine PFSP-SDST').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines (machines non arretées en préparation)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)\nTAP: (teaux d'arret de préparation)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=4)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines PFSP-SDST (machines non arretées en préparation)').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(4):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i == 0 and j == 3:
                    cell.paragraphs[0].runs[0].text = 'TAP'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[2][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[3][i-1]*100,2))+"%"
                elif i != 0 and j == 3:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[4][i-1]*100,2))+"%"
        Figure = plot_Performances([Performances[2],Performances[3],Performances[4]])
        self.doc.add_picture(Figure, width=Inches(5))
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines (machines arretées en préparation)").style = 'Heading2'
        self.Title2+=1
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines PFSP-SDST (machines arretées en préparation)').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[2][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[3][i-1]*100,2))+"%"
        Figure = plot_Performances([Performances[0],Performances[1]])
        self.doc.add_picture(Figure, width=Inches(5))
        self.doc.add_heading("    "+str(self.Title2)+". Temps d'attedre de chaque job entre chaque deux machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TW: (waiting Time) temps d'attedre de chaque job entre chaque deux machines")
        waitting_times = Problem.TW()
        self.table_TW = self.doc.add_table(rows=self.m, cols=self.n+1)
        self.table_TW.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Temps d'attendre de chaque job entre chaque deux machines PFSP-SDST").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m):
            for j in range(self.n+1):
                cell = self.table_TW.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "TW" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)+'/'+str(i)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(waitting_times[0][i-1][j-1]) 
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),0,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))

    def block_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP,block=1)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de BFSP (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("D[i][S[j]]: matrice des dates de fin de chaque job avec son blocage dans chaque machine, aussi début des jobs dans la machine précédente \npour la matrice D l'indexation des machines réel est de 1 à "+str(self.m)+"\n")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrices D,C,F").style = 'Heading2'
        self.Title2+=1
        self.table_D = self.doc.add_table(rows=self.m+2, cols=self.n+1)
        self.table_D.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job avec son blocage dans chaque machine BFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+2):
            for j in range(self.n+1):
                cell = self.table_D.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "D" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[-2][i-1][j-1])

        self.doc.add_paragraph("C[i][S[j]]:dates de fin de chaque job dans chaque machine BFSP\n")
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine BFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])

        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine BFSP \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine BFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])

        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines BFSP ').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[0][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[1][i-1]*100,2))+"%"
        Figure = plot_Performances(Performances)
        self.doc.add_picture(Figure, width=Inches(5)) 
        self.doc.add_heading("    "+str(self.Title2)+". Temps de blocage de chaque job dans chaque machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TB: (Blocking Time) temps de blocage de chaque job dans chaque machines")
        self.table_TB = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_TB.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Temps de blocage de chaque job dans chaque machines BFSP").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_TB.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "TB" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[4][i-1][j-1])      
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),0,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))   

    def prep_block_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP,block=1,prep=1)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de BFSP-SDST (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("D[i][S[j]]: matrice des dates de fin de chaque job avec son blocage dans chaque machine, aussi début des jobs dans la machine précédente \npour la matrice D l'indexation des machines réel est de 1 à "+str(self.m)+" et l'indexation des jobs réel de 1 à "+str(self.n)+"\n")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrices D,C,F").style = 'Heading2'
        self.Title2+=1
        self.table_D = self.doc.add_table(rows=self.m+2, cols=self.n+2)
        self.table_D.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job avec son blocage dans chaque machine BFSP-SDST').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+2):
            for j in range(self.n+2):
                cell = self.table_D.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "D" 
                elif i == 0 and j != 0 and j!=1:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-2]+1)
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'J0'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[-2][i-1][j-1])

        self.doc.add_paragraph("C[i][S[j]]:dates de fin de chaque job dans chaque machine BFSP-SDST\n")
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine BFSP-SDST').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])

        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine BFSP-SDST \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine BFSP-SDST').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])

        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines (machines non arretées en préparation)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)\nTAP: (teaux d'arret de préparation)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=4)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines BFSP-SDST (machines non arretées en préparation)').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(4):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i == 0 and j == 3:
                    cell.paragraphs[0].runs[0].text = 'TAP'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[2][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[3][i-1]*100,2))+"%"
                elif i != 0 and j == 3:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[4][i-1]*100,2))+"%"
        Figure = plot_Performances([Performances[2],Performances[3],Performances[4]])
        self.doc.add_picture(Figure, width=Inches(5))
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines (machines arretées en préparation)").style = 'Heading2'
        self.Title2+=1
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines BFSP-SDST (machines arretées en préparation)').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[2][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[3][i-1]*100,2))+"%"
        Figure = plot_Performances([Performances[0],Performances[1]])
        self.doc.add_picture(Figure, width=Inches(5))
        self.doc.add_heading("    "+str(self.Title2)+". Temps de blocage de chaque job dans chaque machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TB: (Blocking Time) temps de blocage de chaque job dans chaque machines")
        self.table_TB = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_TB.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Temps de blocage de chaque job dans chaque machines BFSP-SDST").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_TB.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "TB" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[6][i-1][j-1])    
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),0,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))

    def noidel_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP,noidel=1)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de NIPFSP (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("C[i][S[j]]: matrice des dates de fin de chaque job dans chaque machine \n")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrice des dates des fins et des débuts de chaque job sur chaque machine").style = 'Heading2'
        self.Title2+=1
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine NIPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])
                                
        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine NIPFSP \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine NIPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines NIPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[0][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[1][i-1]*100,2))+"%"
        Figure = plot_Performances(Performances)
        self.doc.add_picture(Figure, width=Inches(5)) 
        self.doc.add_heading("    "+str(self.Title2)+". Temps d'attedre de chaque job entre chaque deux machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TW: (waiting Time) temps d'attedre de chaque job entre chaque deux machines")
        self.table_TW = self.doc.add_table(rows=self.m, cols=self.n+1)
        self.table_TW.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Temps d'attendre de chaque job entre chaque deux machines NIPFSP").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m):
            for j in range(self.n+1):
                cell = self.table_TW.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "TW" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)+'/'+str(i)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[2][i-1][j-1])
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),0,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))

    def nowait_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP,nowait=1)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de NWPFSP (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("C[i][S[j]]: matrice des dates de fin de chaque job dans chaque machine \n")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrice des dates des fins et des débuts de chaque job sur chaque machine").style = 'Heading2'
        self.Title2+=1
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine NWPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])
                                
        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine NWPFSP \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine NWPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines NWPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[0][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[1][i-1]*100,2))+"%"
        Figure = plot_Performances(Performances)
        self.doc.add_picture(Figure, width=Inches(5)) 
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),1,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))

    def save_doc(self):
        self.doc.save(self.file_path)
        print(f"Word document created at: {self.file_path}")

def get_save_path():
    root = Tk()
    root.withdraw()  # Hide the main window

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")],
        title="Save Word Document As"
    )

    root.destroy()  # Close the Tkinter window

    return file_path

def plot_Performances(TFR_TAR_plt):
    categories = np.arange(len(TFR_TAR_plt[0]))
    bar_width = 0.25
    names=["TFR","TAR","TAP"]
    fig, ax = plt.subplots()
    plt.figure(figsize=(20, 20))
    for i, t_list in enumerate(TFR_TAR_plt):
        bars = ax.bar(categories + i * bar_width, t_list, label=names[i], width=bar_width, alpha=0.7)
        for bar, prob in zip(bars, t_list):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5, "", ha='center', va='bottom')

    ax.set_xlabel('Machines')
    ax.set_ylabel('Performances(%)')
    ax.set_title("temps de fonctionnement réel(TFR) et Temps d'arret réel/préparation (TAR/TAP)")
    ax.set_xticks(categories)
    legend = ax.legend(loc='lower right')
    legend.set_draggable(True)
    image_stream = BytesIO()
    fig.savefig(image_stream, format='png')
    image_stream.seek(0)
    return image_stream

def gantt_plot(tasks,nowait,name):

    fig, ax = plt.subplots()
    # Configuration de l'axe Y
    ax.set_yticks(range(len(tasks)))
    ax.set_yticklabels(tasks.index)

    # Flattening the lists of lists
    all_start_times = [time for times_list in tasks['Start'] for times in times_list for time in times]
    all_end_times = [time for times_list in tasks['End'] for times in times_list for time in times]

    # Configuration de l'axe X
    ax.set_xticks(np.arange(min(all_start_times), max(all_end_times) + 1, step=1))
    ax.set_xlim(min(all_start_times), max(all_end_times) + 1)

    # Tracé des tâches
    for i, task in enumerate(tasks.itertuples()):
        for start_times, end_times, color, label in zip(task.Start, task.End, task.Color, task.Label):
            for start, end, text in zip(start_times, end_times, label):
                # Check if the color is a valid color name or hexadecimal code
                try:
                    mcolors.to_rgba(color)
                except ValueError:
                    color = 'blue'  # Default to blue if an invalid color is provided
                ax.barh(i, end - start, left=start, color=color, alpha=0.7)
                text_x = start + (end - start) / 2
                text_y = i
                ax.text(text_x, text_y, label, ha='center', va='center', color='white', fontweight='bold')

                # Add an arrow at the end of each job except for the first one
                if i > 0 and nowait==1:
                    arrow_start = end
                    arrow_end = end  # The arrow origin is at the bottom of the job
                    ax.arrow(arrow_start, text_y-0.4, 0, -0.1, head_width=0.2, head_length=0.1, fc=color, ec=color)

    # Configuration des étiquettes
    ax.set_xlabel('Axe de temps')
    ax.set_title('Diagramme de Gantt('+name+')')
    plt.grid()
    image_stream = BytesIO()
    fig.savefig(image_stream, format='png')
    image_stream.seek(0)
    return image_stream




class Tableaux(ct.CTk):
    def __init__(self,endJ,startJ,endP,startP,TFR_TAR,TW,TWA,TWAN,TBF,tasks,seq,prep,block,nowait,noidel,TFT,TT,Tm,use_TT):
        super().__init__()
        self.title("resultat de Gantt")
        self.endJ = endJ
        self.startJ = startJ
        self.endP = endP
        self.startP = startP
        self.T_FA_R = TFR_TAR
        self.TW = TW
        self.TWA = TWA
        self.TWAN = TWAN
        self.tasks = tasks
        self.seq = seq
        self.prep = prep
        self.block = block
        self.TBF = TBF
        self.cas = 1
        self.nowait = nowait
        self.noidel = noidel
        self.TFT = TFT
        self.TT = TT
        self.Tm = Tm
        self.use_TT = use_TT
        for i in range(len(self.T_FA_R)):
            for j in range(len(self.T_FA_R[0])):
                var=round(self.T_FA_R[i][j]*100,2)
                self.T_FA_R[i][j]=str(var)+"%"
        self.tab_control = ct.CTkTabview(self,width=1400,height=800)

        if self.prep == 1 and self.block == 0:
            self.tab_full_names=["Diagramme de Gaintt"
                                ,"Dates de fin des Jobs"
                                ,"Dates de début des Jobs"
                                ,"Dates de fin des temps de preparation"
                                ,"Dates de début des temps de preparation"
                                ,"temps de fonctionnement  réel \n et temps d'arret réel et temps d'arret de préparation \n cas 1 : preparation = arret"
                                ,"Temps d'attente des Jobs"
                                ,"Temps d'attente des Jobs à cause de preparation"
                                ,"Temps d'attente des Jobs à cause de fonctionnement"]
            self.tab_names=["Gantt","FinJ","DébutJ","FinP","DébutP","T_FA_R","TA","TAP","TAF"]
            for name in self.tab_names:
                self.tab_control.add(name)
                self.tab_control.pack(side=tk.LEFT, fill='y')
            cas1_btn = ct.CTkButton(master=self.tab_control.tab("T_FA_R"), text= "Preparation = Arret", command=self.cas1_fun)
            cas1_btn.grid(row=1,column=0, padx=40, pady=20)
            cas2_btn = ct.CTkButton(master=self.tab_control.tab("T_FA_R"), text= "Preparation = Travail", command=self.cas2_fun)
            cas2_btn.grid(row=2,column=0, padx=40, pady=20)
            self.plot_gantt()
            self.create_endJ()
            self.create_startJ()
            self.create_endP()
            self.create_startP()
            self.create_T_FA_R()
            self.create_TW()
            self.create_TWA()
            self.create_TWAN()
        elif self.prep == 0 and self.block == 0:
            self.T_FA_R = transpose_matrix(self.T_FA_R)
            self.tab_full_names=["Diagramme de Gaintt"
                                ,"Dates de fin des Jobs"
                                ,"Dates de début des Jobs"
                                ,"temps de fonctionnement/arret réel"
                                ,"Temps d'attente des Jobs"]
            self.tab_names=["Gantt","FinJ","DébutJ","T_FA_R","TA"]
            for name in self.tab_names:
                self.tab_control.add(name)
                self.tab_control.pack(side=tk.LEFT, fill='y')
            self.plot_gantt()
            self.create_endJ()
            self.create_startJ()
            self.create_T_FA_R()
            self.create_TW()
        elif self.prep == 0 and self.block == 1:
            self.T_FA_R = transpose_matrix(self.T_FA_R)
            self.tab_full_names=["Diagramme de Gaintt"
                                ,"Dates de fin des Jobs"
                                ,"Dates de début des Jobs"
                                ,"temps de fonctionnement/arret réel"
                                ,"Temps de blocage des Jobs"]
            self.tab_names=["Gantt","FinJ","DébutJ","T_FA_R","TB"]
            for name in self.tab_names:
                self.tab_control.add(name)
                self.tab_control.pack(side=tk.LEFT, fill='y')
            self.plot_gantt()
            self.create_endJ()
            self.create_startJ()
            self.create_T_FA_R()
            self.create_TW()
        elif self.prep == 1 and self.block == 1:
            self.tab_full_names=["Diagramme de Gaintt"
                                ,"Dates de fin des Jobs"
                                ,"Dates de début des Jobs"
                                ,"Dates de fin des temps de preparation"
                                ,"Dates de début des temps de preparation"
                                ,"temps de fonctionnement  réel \n et temps d'arret réel et temps d'arret de préparation \n cas 1 : preparation = arret"
                                ,"Temps de blocage des Jobs"
                                ,"Temps de blocage des Jobs à cause de preparation"
                                ,"Temps de blocage des Jobs à cause de blocage"
                                ,"Temps de blocage des Jobs à cause de fonctionnement"]
            self.tab_names=["Gantt","FinJ","DébutJ","FinP","DébutP","T_FA_R","TB","TBP","TBB","TBF"]
            for name in self.tab_names:
                self.tab_control.add(name)
                self.tab_control.pack(side=tk.LEFT, fill='y')
            cas1_btn = ct.CTkButton(master=self.tab_control.tab("T_FA_R"), text= "Preparation = Arret", command=self.cas1_fun)
            cas1_btn.grid(row=1,column=0, padx=40, pady=20)
            cas2_btn = ct.CTkButton(master=self.tab_control.tab("T_FA_R"), text= "Preparation = Travail", command=self.cas2_fun)
            cas2_btn.grid(row=2,column=0, padx=40, pady=20)
            self.plot_gantt()
            self.create_endJ()
            self.create_startJ()
            self.create_endP()
            self.create_startP()
            self.create_T_FA_R()
            self.create_TW()
            self.create_TWA()
            self.create_TWAN()
            self.create_TBF()
        if self.use_TT==1:
            self.tab_control.add('T')
            self.tab_control.pack(side=tk.LEFT, fill='y')
            self.create_Tm()

    def create_endJ(self):
        self.create_regular_tab(self.endJ,self.tab_names[1],self.tab_full_names[1])

    def create_startJ(self):
        self.create_regular_tab(self.startJ,self.tab_names[2],self.tab_full_names[2])

    def create_endP(self):
        self.create_regular_tab(self.endP,self.tab_names[3],self.tab_full_names[3])

    def create_startP(self):
        self.create_regular_tab(self.startP,self.tab_names[4],self.tab_full_names[4])

    def create_T_FA_R(self):
        if self.prep == 1:
            TFAR = transpose_matrix([self.T_FA_R[0],self.T_FA_R[1],self.T_FA_R[4]])
            print(TFAR,"TFAR")
            self.create_regular_tab(TFAR,self.tab_names[5],self.tab_full_names[5])
        else:
            self.create_regular_tab(self.T_FA_R,self.tab_names[3],self.tab_full_names[3])

    def create_TW(self):
        if self.prep ==1:
            self.create_regular_tab(self.TW,self.tab_names[6],self.tab_full_names[6])
        else:
            self.create_regular_tab(self.TW,self.tab_names[4],self.tab_full_names[4])

    def create_TWA(self):
        self.create_regular_tab(self.TWA,self.tab_names[7],self.tab_full_names[7])

    def create_TWAN(self):
        self.create_regular_tab(self.TWAN,self.tab_names[8],self.tab_full_names[8])
    
    def create_TBF(self):
        self.create_regular_tab(self.TBF,self.tab_names[9],self.tab_full_names[9])

    def create_Tm(self):
        self.create_regular_tab([self.Tm],"T","total flow time (TFT="+str(self.TFT)+ ") / Tardiness of jobs (Tj) / Total tardness (TT="+str(self.TT)+")")

    def create_regular_tab(self , tab, name, full_name):
        n , m = len(tab[0]) , len(tab)
        table_label = ct.CTkLabel(self.tab_control.tab(name),text=full_name)
        tree = ttk.Treeview(self.tab_control.tab(name), columns=list(range(n+1)),show='headings')
        if name == "FinJ" or name =="DébutJ" or name =="FinP" or name =="DébutP" or name == "T":
            table_label.pack(side=tk.TOP)
            for j in range(n+1):
                if j!=0:
                    tree.heading(j, text=f'Job {self.seq[j-1]}')
                else:
                    tree.heading(j, text="Machines")
                tree.column(j, width=80)

            for i in range(m):
                if type(tab[i]) != list:
                    if name == 'T':
                        tree.insert("", "end", values=['M'+str(len(self.endJ)-1)]+tab[i])
                    else:
                        tree.insert("", "end", values=['M'+str(i)]+tab[i].tolist())
                else:
                    if name == 'T':
                        tree.insert("", "end", values=['M'+str(len(self.endJ)-1)]+tab[i])
                    else:
                        tree.insert("", "end", values=['M'+str(i)]+tab[i])
            tree.pack()
        elif name == "T_FA_R":
            self.table_labelt = ct.CTkLabel(self.tab_control.tab(name),text=full_name)
            self.treet = ttk.Treeview(self.tab_control.tab(name), columns=list(range(n+1)),show='headings')
            table_label.grid(row=0,column=0,columnspan=5)
            self.treet.heading(0, text="Machines")
            self.treet.column(0, width=80)
            self.treet.heading(1, text="TFR")
            self.treet.column(1, width=80)
            self.treet.heading(2, text="TAR")
            self.treet.column(2, width=80)
            TFR_TAR_plt = [[float(t.strip('%')) for t in lst] for lst in tab]
            TFR_TAR_plt = np.array(TFR_TAR_plt).T.tolist()
            categories = np.arange(len(TFR_TAR_plt[0]))
            bar_width = 0.25
            names=["TFR","TAR","TAP"]
            fig, ax = plt.subplots()
            plt.figure(figsize=(20, 20))
            for i, t_list in enumerate(TFR_TAR_plt):
                print(names,i,"tttt")
                bars = ax.bar(categories + i * bar_width, t_list, label=names[i], width=bar_width, alpha=0.7)
                for bar, prob in zip(bars, t_list):
                    ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5, f'{prob:.2f}%', ha='center', va='bottom')
                if i ==2:
                    self.treet.heading(3, text="TPR")
                    self.treet.column(3, width=80)
            ax.set_xlabel('Machines')
            ax.set_ylabel('Performances(%)')
            ax.set_title("temps de fonctionnement réel(TFR) et Temps d'arret réel/préparation (TAR/TAP)")
            ax.set_xticks(categories)
            legend = ax.legend(loc='lower right')
            legend.set_draggable(True)
            self.canvas_T = FigureCanvasTkAgg(fig, master=self.tab_control.tab(name))
            self.canvas_T_widget = self.canvas_T.get_tk_widget()
            self.canvas_T_widget.place(x=500,y=200)
            self.canvas_T.draw()
            self.canvas_T.flush_events()
            self.treet.grid(row=3,column=0,rowspan=3)
            for i in range(m):
                self.treet.insert("", "end", values=['M'+str(i)]+tab[i])
        elif name =="TA" or name =="TAP" or name =="TAF":
            table_label.pack(side=tk.TOP)
            for j in range(n+1):
                if j!=0:
                    if self.nowait==1:
                        tree.heading(j, text=f'Job'+ str(self.seq[j-1])+"/"+str(self.seq[j]))
                    else:
                        tree.heading(j, text=f'Job {self.seq[j-1]}')
                else:
                    tree.heading(j, text="Machines")
                tree.column(j, width=80)

            for i in range(m):
                if self.nowait==1:
                    if type(tab[i]) != list:
                        tree.insert("", "end", values=['M'+str(i)]+tab[i].tolist())
                    else:
                        tree.insert("", "end", values=['M'+str(i)]+tab[i]) 
                else:
                    if type(tab[i]) != list:
                        tree.insert("", "end", values=['M'+str(i)+"/"+str(i+1)]+tab[i].tolist())
                    else:
                        tree.insert("", "end", values=['M'+str(i)+"/"+str(i+1)]+tab[i])
            tab = transpose_matrix(tab)
            sums=[sum(lst) for lst in tab]
            tree.insert("", "end", values=["somme"]+sums)
            tree.pack()
        elif name == "TB" or name == "TBP" or name == "TBF" or name == "TBB":
            table_label.pack(side=tk.TOP)
            for j in range(n+1):
                if j!=0:
                    tree.heading(j, text=f'Job {self.seq[j-1]}')
                else:
                    tree.heading(j, text="Machines")
                tree.column(j, width=80)

            for i in range(m):
                if type(tab[i]) != list:
                    tree.insert("", "end", values=['M'+str(i)]+tab[i].tolist())
                else:
                    tree.insert("", "end", values=['M'+str(i)]+tab[i])
            tab = transpose_matrix(tab)
            sums=[sum(lst) for lst in tab]
            tree.insert("", "end", values=["somme"]+sums)
            tree.pack()
    
    def cas1_fun(self):
        if self.cas == 2:
            self.table_labelt.destroy()
            self.treet.destroy()
            self.create_T_FA_R()
            self.cas = 1
        else:
            pass
        
    def cas2_fun(self):
        if self.cas == 1:
            self.table_labelt.destroy()
            self.treet.destroy()
            self.create_regular_tab(transpose_matrix([self.T_FA_R[2],self.T_FA_R[3]]),self.tab_names[5],"temps de fonctionnement  réel \n et temps d'arret réel et temps d'arret de préparation \n cas 2 : preparation = travail")
            self.cas = 2
        else:
            pass
    
    def plot_gantt(self):
        tasks = self.tasks

        fig, ax = plt.subplots()
        # Configuration de l'axe Y
        ax.set_yticks(range(len(tasks)))
        ax.set_yticklabels(tasks.index)

        # Flattening the lists of lists
        all_start_times = [time for times_list in tasks['Start'] for times in times_list for time in times]
        all_end_times = [time for times_list in tasks['End'] for times in times_list for time in times]

        # Configuration de l'axe X
        ax.set_xticks(np.arange(min(all_start_times), max(all_end_times) + 1, step=1))
        ax.set_xlim(min(all_start_times), max(all_end_times) + 1)

        self.canvas_gantt = FigureCanvasTkAgg(fig, master=self.tab_control.tab("Gantt"))
        self.canvas_gantt_widget = self.canvas_gantt.get_tk_widget()
        self.canvas_gantt_widget.pack(side=tk.TOP, fill=tk.BOTH, expand=1)

        # Tracé des tâches
        for i, task in enumerate(tasks.itertuples()):
            for start_times, end_times, color, label in zip(task.Start, task.End, task.Color, task.Label):
                for start, end, text in zip(start_times, end_times, label):
                    # Check if the color is a valid color name or hexadecimal code
                    try:
                        mcolors.to_rgba(color)
                    except ValueError:
                        color = 'blue'  # Default to blue if an invalid color is provided
                    ax.barh(i, end - start, left=start, color=color, alpha=0.7)
                    text_x = start + (end - start) / 2
                    text_y = i
                    ax.text(text_x, text_y, label, ha='center', va='center', color='white', fontweight='bold')

                    # Add an arrow at the end of each job except for the first one
                    if i > 0 and self.nowait==1:
                        arrow_start = end
                        arrow_end = end  # The arrow origin is at the bottom of the job
                        ax.arrow(arrow_start, text_y-0.4, 0, -0.1, head_width=0.2, head_length=0.1, fc=color, ec=color)

        # Configuration des étiquettes
        ax.set_xlabel('Axe de temps')
        if self.use_TT==1:
            ax.set_title('Diagramme de Gantt \n Total flow tmie (TFT) = '+str(self.TFT)+"\n Cmax = "+str(self.endJ[-1][len(self.endJ[0])-1])+"\n Total Tardiness (TT) = "+str(self.TT))
        else:
            ax.set_title('Diagramme de Gantt \n Total flow tmie (TFT) = '+str(self.TFT)+"\n Cmax = "+str(self.endJ[-1][len(self.endJ[0])-1]))
        plt.grid()

        self.canvas_gantt.draw()
        self.canvas_gantt.flush_events()


















# cas possible de (prep,block,noidel,nowait) = (0,0,0,0)/(1,0,0,0)/(0,1,0,0)/(1,1,0,0)/(0,0,1,0)/(0,0,0,1)
# Définition de la classe d'ordonnancement qui sert à traiter les problemes de flowshop
class ordonnancement () :
    def __init__(self,O=[],F=[],OP=[],FP=[],M=[[]],S=[],Mprep=[[]],prep=0,block=0,noidel=0,nowait=0,dj=[],TT=0,sim=0,do_report=0): 
        self.O = O           # O: dates des fins des jobs (out of job)
        self.M = M           # M: matrice des temps de processus 
        self.S = S           # S: séquence des jobs (indexage début de 0)
        self.F = F           # F: dates des debuts des jobs (first of job)
        self.OP = OP         # OP : dates des fins des temps de préparation (out of preparation time)
        self.FP = FP         # FP : dates des debuts des temps de préparation (firt of preparation time)       
        self.Mprep = Mprep   # Mprep : matrice des temps de preparations
        self.prep = prep     # si = 0 on n'utiluse pas les temps de preparation si non utiluse les
        self.block = block   # si = 0 on n'utiluse pas la condition de blockage si non utiluse la
        self.noidel = noidel # si = 0 on n'utiluse pas la condition de noidel si non utiluse la
        self.nowait = nowait # si = 0 on n'utiluse pas la condition de nowait si non utiluse la
        self.dj = dj         # dj : list des delais des jobs 
        self.TT = TT         # si = 0 on n'utiluse pas les delais si non utiluse les
        self.sim = sim       # si = 0 on return un dataframe dans gatt si on on retourn des listre pour les simule
        self.do_report = do_report 
        self.report = {}
    # Méthode pour traiter le probleme flowshop
    def traite (self):
        # initialisation de variables 
        self.O,self.OP,self.F,self.FP=[],[],[],[]
        # avec condition de non arret
        if self.noidel == 1:
            return noidel_probleme_solving(self.M,self.S,self.dj,self.TT,self.do_report)
        # avec condition de non attendre
        if self.nowait == 1:
            return nowait_probleme_solving(self.M,self.S,self.dj,self.TT,self.do_report)
        # sans preparation et sans blockage 
        if self.prep == 0 and self.block == 0:
            char = "Sans préparation et sans aucune condition \n" 
            # initialisation
            for i in range(len(self.M)):
                self.O.append([0.0]*len(self.M[0]))
                self.F.append([0.0]*len(self.M[0]))
            # equation 1
            self.O[0][0] = self.M[0][self.S[0]]
            char+='équation 1: pour i = 0 et j = 0\n'
            char+='  C[0][J'+str(self.S[0])+'] = P[0][J'+str(self.S[0])+'] = '+str(self.O[0][0])+'\n'
            # equation 2
            char+='équation 2: pour i = 0,...,'+str(len(self.M)-1)+' et j = 0'+'\n'
            for i in range(1,len(self.M)):
                char+='  C['+str(i)+'][J'+str(self.S[0])+'] = C['+str(i-1)+'][J'+str(self.S[0])+'] + P['+str(i)+'][J'+str(self.S[0])+'] = '+str(self.O[i-1][0])+" + "+str(self.M[i][self.S[0]])+" = "+str(self.O[i-1][0]+self.M[i][self.S[0]])+'\n'
                self.O[i][0] = self.O[i-1][0]+self.M[i][self.S[0]]
                self.F[i][0] = self.O[i-1][0]
            # equation 3
            char+='équation 3: pour i = 0 et j = 0,...,'+str(len(self.M[0])-1)+'\n'
            for j in range(1,len(self.S)):
                char+='  C[0][J'+str(self.S[j])+'] = C[0][J'+str(self.S[j-1])+'] + P[0][J'+str(self.S[j])+'] = '+str(self.O[0][j-1])+" + "+str(self.M[0][self.S[j]])+" = "+str(self.O[0][j-1]+self.M[0][self.S[j]])+'\n'
                self.O[0][j] = self.O[0][j-1]+self.M[0][self.S[j]]
                self.F[0][j] = self.O[0][j-1]
            # equation 4
            char+='équation 4: pour i = 1,...,'+str(len(self.M)-1)+' et j = 1,...,'+str(len(self.M[0])-1)+'\n'
            for i in range(1,len(self.M)):
                for j in range(1,len(self.S)):
                    char+='  C['+str(i)+'][J'+str(self.S[j])+'] = max( C['+str(i-1)+'][J'+str(self.S[j])+'] , C['+str(i)+'][J'+str(self.S[j-1])+'] ) + P['+str(i)+'][J'+str(self.S[j])+'] = max('+str(self.O[i-1][j])+','+str(self.O[i][j-1])+')'+' '+str(self.M[i][self.S[j]])+' = '+str(max(self.O[i-1][j],self.O[i][j-1]))+'+'+str(self.M[i][self.S[j]])+' = '+str(max(self.O[i-1][j],self.O[i][j-1])+self.M[i][self.S[j]])+'\n'
                    self.O[i][j] = max(self.O[i-1][j],self.O[i][j-1])+self.M[i][self.S[j]]
                    self.F[i][j] = max(self.O[i-1][j],self.O[i][j-1])
            if self.TT == 1:
                char2=""
                char3="Tardiness T:\n"
                char+="\n"+'Total flow time TFT = '
                TFT,Tm,TT=0.0,[],0.0
                #TFT : total flow time
                #Tm :temps de retard d'un job "tardiness"
                #TT : total tardiness
                for j in range(len(self.M[0])):
                    char+= '  C['+str(len(self.M)-1)+'][J'+str(self.S[j])+']'
                    char2+= str(self.O[-1][j])
                    char3+= '  T[J'+str(self.S[j])+'] = '+'max( C['+str(len(self.M)-1)+'][J'+str(self.S[j])+'] - d[J'+str(self.S[j])+'] , 0 ) = '+'max( '+str(self.O[-1][j])+' - '+str(self.dj[self.S[j]])+' , 0 ) = '+str(max(self.O[-1][j]-self.dj[self.S[j]],0))+'\n'
                    if j!=len(self.M[0])-1:
                        char+=" + "
                        char2+=" + "
                    else:
                        char+=" = "
                        char2+=" = "
                    TFT+=self.O[-1][j]
                    Tm.append(max(self.O[-1][j]-self.dj[self.S[j]],0))
                char+="\n"+char2
                char+=str(TFT)+'\n'
                char3+= "Tardiness TT = ∑ T[S[j]] : j = 0,...,"+str(len(self.M[0])-1)+"\n  TT = "+str(sum(Tm))
                char+="\n"+char3
                TT=sum(Tm)
                # renvoi des dates des fins et des debuts de chaque job
                if self.do_report == 0:
                    return [self.O,self.F,TFT,Tm,TT]
                else:
                    return [self.O,self.F,TFT,Tm,TT,char]
            else:
                char2=""
                char+="\n"+'Total flow time TFT = '
                TFT=0.0
                for j in range(len(self.M[0])):
                    char+= 'C['+str(len(self.M)-1)+'][J'+str(self.S[j])+']'
                    char2+= str(self.O[-1][j])
                    if j!=len(self.M[0])-1:
                        char+=" + "
                        char2+=" + "
                    else:
                        char+=" = "
                        char2+=" = "
                    TFT+=self.O[-1][j]
                char+="\n"+char2
                char+=str(TFT)
                print(char)
                if self.do_report==0:
                    return [self.O,self.F,TFT]
                else:
                    return [self.O,self.F,TFT,char]
        # avec preparaion et sans blockage
        elif self.prep == 1 and self.block == 0:
            char = "Avec préparation et sans aucune condition \n"
            # initialisation
            for i in range(len(self.M)):
                self.O.append([0.0]*len(self.M[0]))
                self.F.append([0.0]*len(self.M[0]))
                self.OP.append([0.0]*len(self.M[0]))
                self.FP.append([0.0]*len(self.M[0]))
            # equation 1
            char+='équation 1: pour i = 0 et j = 0\n'
            char+='  C[0][J'+str(self.S[0])+'] = P[0][J'+str(self.S[0])+'] + M[0][J'+str(self.S[0])+'][J'+str(self.S[0])+'] = '+str(self.M[0][self.S[0]])+'+'+str(self.Mprep[0][self.S[0]][self.S[0]])+' = '+str(self.M[0][self.S[0]]+self.Mprep[0][self.S[0]][self.S[0]])+'\n'
            self.O[0][0] = self.M[0][self.S[0]]+self.Mprep[0][self.S[0]][self.S[0]]
            self.F[0][0] = self.Mprep[0][self.S[0]][self.S[0]]
            self.OP[0][0] = self.F[0][0]
            # equation 2
            char+='équation 2: pour i = 1,...,'+str(len(self.M)-1)+' et j = 0\n'
            for i in range(1,len(self.M)):         
                char+='  C['+str(i)+'][J'+str(self.S[0])+'] = max( C['+str(i-1)+'][J'+str(self.S[0])+'] , M['+str(i)+'][J'+str(self.S[0])+'][J'+str(self.S[0])+'] ) + P['+str(i)+'][J'+str(self.S[0])+'] = max( '+str(self.O[i-1][0])+' , '+str(self.Mprep[0][self.S[0]][self.S[0]])+') + '+str(self.M[i][self.S[0]])+' = '+str(max( self.O[i-1][0] , self.Mprep[i][self.S[0]][self.S[0]] )+self.M[i][self.S[0]])+'\n'
                self.O[i][0] = max( self.O[i-1][0] , self.Mprep[i][self.S[0]][self.S[0]] )+self.M[i][self.S[0]]
                self.F[i][0] = max( self.O[i-1][0] , self.Mprep[i][self.S[0]][self.S[0]] )
                self.OP[i][0] = self.Mprep[i][self.S[0]][self.S[0]]
            # equation 3
            char+='équation 3: pour i = 0 et j = 1,...,'+str(len(self.M[0])-1)+'\n'
            for j in range(1,len(self.S)):
                char+='  C[0][J'+str(self.S[j])+'] = C[0][J'+str(self.S[j-1])+'] + P[0][J'+str(self.S[j-1])+'] + M[0][J'+str(self.S[j-1])+'][J'+str(self.S[j])+'] = '+str(self.O[0][j-1])+'+'+str(self.M[0][self.S[j]])+'+'+str(self.Mprep[0][self.S[j-1]][self.S[j]])+' = '+str(self.O[0][j-1]+self.M[0][self.S[j]]+self.Mprep[0][self.S[j-1]][self.S[j]])+'\n'
                self.O[0][j] = self.O[0][j-1]+self.M[0][self.S[j]]+self.Mprep[0][self.S[j-1]][self.S[j]]
                self.F[0][j] = self.O[0][j-1]+self.Mprep[0][self.S[j-1]][self.S[j]]
                self.OP[0][j] = self.F[0][j]
                self.FP[0][j] = self.O[0][j-1]
            # equation 4
            char+='équation 4: pour i = 1,...,'+str(len(self.M)-1)+'j = 1,...,'+str(len(self.M[0])-1)+'\n'
            for i in range(1,len(self.M)):
                for j in range(1,len(self.S)):
                    char+='  C['+str(i)+'][J'+str(self.S[j])+'] = max( C['+str(i-1)+'][J'+str(self.S[j])+'] , C['+str(i)+'][J'+str(self.S[j-1])+'] + M['+str(i)+'][J'+str(self.S[j-1])+'][J'+str(self.S[j])+'] ) + P['+str(i)+'][J'+str(self.S[j])+'] = max( '+str(self.O[i-1][j])+' , '+str(self.O[i][j-1])+'+'+str(self.Mprep[i][self.S[j-1]][self.S[j]])+') + '+str(self.M[i][self.S[0]])+' = '+str(max(self.O[i-1][j],self.O[i][j-1]+self.Mprep[i][self.S[j-1]][self.S[j]])+self.M[i][self.S[j]])+'\n'
                    self.O[i][j] = max(self.O[i-1][j],self.O[i][j-1]+self.Mprep[i][self.S[j-1]][self.S[j]])+self.M[i][self.S[j]]
                    self.F[i][j] = max(self.O[i-1][j],self.O[i][j-1]+self.Mprep[i][self.S[j-1]][self.S[j]])
                    self.OP[i][j] = self.O[i][j-1]+self.Mprep[i][self.S[j-1]][self.S[j]]
                    self.FP[i][j]= self.O[i][j-1]
            if self.TT == 1:
                char2=""
                char3="Tardiness T:\n"
                char+="\n"+'Total flow time TFT = '
                TFT,Tm,TT=0.0,[],0.0
                #TFT : total flow time
                #Tm :temps de retard d'un job "tardiness"
                #TT : total tardiness
                for j in range(len(self.M[0])):
                    char+= '  C['+str(len(self.M)-1)+'][J'+str(self.S[j])+']'
                    char2+= str(self.O[-1][j])
                    char3+= '  T[J'+str(self.S[j])+'] = '+'max( C['+str(len(self.M)-1)+'][J'+str(self.S[j])+'] - d[J'+str(self.S[j])+'] , 0 ) = '+'max( '+str(self.O[-1][j])+' - '+str(self.dj[self.S[j]])+' , 0 ) = '+str(max(self.O[-1][j]-self.dj[self.S[j]],0))+'\n'
                    if j!=len(self.M[0])-1:
                        char+=" + "
                        char2+=" + "
                    else:
                        char+=" = "
                        char2+=" = "
                    TFT+=self.O[-1][j]
                    Tm.append(max(self.O[-1][j]-self.dj[self.S[j]],0))
                char+="\n"+char2
                char+=str(TFT)+'\n'
                char3+= "Tardiness TT = ∑ T[S[j]] : j = 0,...,"+str(len(self.M[0])-1)+"\n  TT = "+str(sum(Tm))
                char+="\n"+char3
                TT=sum(Tm)
                # renvoi des dates des fins et des debuts de chaque job et temp de preparation
                if self.do_report==0:
                    return [self.O,self.F,self.OP,self.FP,TFT,Tm,TT]
                else:
                    return [self.O,self.F,self.OP,self.FP,TFT,Tm,TT,char]
            else:
                char2=""
                char+="\n"+'Total flow time TFT = '
                TFT=0.0
                for j in range(len(self.M[0])):
                    char+= '  C['+str(len(self.M)-1)+'][J'+str(self.S[j])+']'
                    char2+= str(self.O[-1][j])
                    if j!=len(self.M[0])-1:
                        char+=" + "
                        char2+=" + "
                    else:
                        char+=" = "
                        char2+=" = "
                    TFT+=self.O[-1][j]
                char+="\n"+char2
                char+=str(TFT)
                print(char)
                if self.do_report==0:
                    return [self.O,self.F,self.OP,self.FP,TFT]
                else:
                    return [self.O,self.F,self.OP,self.FP,TFT,char]
        # sans preparation et avec blockage 
        elif self.prep == 0 and self.block == 1:
            # self.D est la matrice des fins des jobs avec leur blockage qui nous donne plusieur donnée 
            # initialisation et equation 1
            char = "Sans préparation et avec blocage \n on Ajout une machine vertuel pour calculer la matrice D\n indéxation des machines réeel de 1 à"+str(len(self.M))+'\n'
            self.D = np.array([[0.0]*len(self.M[0])]*(len(self.M)+1))
            char+='équation 1: pour i = 0 et j = 0\n'
            char+='  D[0][J'+str(self.S[0])+'] = 0\n'
            # equation 2
            char+='équation 2: pour i = 1,...,'+str(len(self.M)-1)+' et j = 0\n'
            for i in range(1,len(self.M)):
                char+='  D['+str(i)+'][J'+str(self.S[0])+'] = D['+str(i-1)+'][J'+str(self.S[0])+'] + P['+str(i-1)+'][J'+str(self.S[0])+'] = '+str(self.D[i-1][0])+' + '+str(self.M[i-1][self.S[0]])+' = '+str(self.D[i-1][0]+self.M[i-1][self.S[0]])+'\n'
                self.D[i][0]=self.D[i-1][0]+self.M[i-1][self.S[0]]
            # equation 5
            char+='équation 5: pour i = '+str(len(self.M)-1)+' et j = 0\n'
            char+='  D['+str(len(self.M))+'][J'+str(self.S[0])+'] = D['+str(len(self.M)-1)+'][J'+str(self.S[0])+'] + P['+str(len(self.M)-1)+'][J'+str(self.S[0])+'] = '+str(self.D[len(self.M)-1][0])+' + '+str(self.M[len(self.M)-1][self.S[0]])+' = '+str(self.D[len(self.M)-1][0]+self.M[len(self.M)-1][self.S[0]])+'\n'
            self.D[len(self.M)][0]=self.D[len(self.M)-1][0]+self.M[len(self.M)-1][self.S[0]]
            # equation 2,3,4
            for j in range(1,len(self.D[0])):
                for i in range(len(self.D)):
                    if i == 0:# equation 3
                        char+='équation 3: pour i = 0 et j = '+str(j)+'\n'
                        char+='  D[0][J'+str(self.S[j])+'] = D[1][J'+str(self.S[j-1])+'] = '+str(self.D[1][j-1])+'\n'
                        self.D[0][j]=self.D[1][j-1]
                    elif i!=0 and i!=len(self.M):# equation 4
                        char+='équation 4: pour i = '+str(i)+' et j = '+str(j)+'\n'
                        char+='  D['+str(i)+'][J'+str(self.S[j])+'] = max( D['+str(i-1)+'][J'+str(self.S[j])+'] + P['+str(i-1)+'][J'+str(max(self.D[i-1][j]+self.M[i-1][self.S[j]],self.D[i+1][j-1]))+'\n'
                        self.D[i][j]=max(self.D[i-1][j]+self.M[i-1][self.S[j]],self.D[i+1][j-1])
                    else:# equation 5
                        char+='équation 5: pour i = '+str(len(self.M)-1)+' et j = '+str(j)+'\n'
                        char+='  D['+str(i)+'][J'+str(self.S[j])+'] = D['+str(len(self.M)-1)+'][J'+str(self.S[j])+'] + P['+str(len(self.M)-1)+'][J'+str(self.S[j])+'] = '+str(self.D[len(self.M)-1][j])+' + '+str(self.M[len(self.M)-1][self.S[j]])+' = '+str(self.D[len(self.M)-1][j]+self.M[len(self.M)-1][self.S[j]])+'\n'
                        self.D[i][j]=self.D[len(self.M)-1][j]+self.M[len(self.M)-1][self.S[j]]
            self.F = np.array([[0.0]*len(self.M[0])]*len(self.M)) # dates de debuts des jobs 
            self.DN= np.array([[0.0]*len(self.M[0])]*len(self.M)) # nouveau matrice D
            self.TB= np.array([[0.0]*len(self.M[0])]*len(self.M)) # temps de blockage
            self.O = np.array([[0.0]*len(self.M[0])]*len(self.M)) # dates des fins des jobs 
            self.TBO = np.array([[0.0]*len(self.M[0])]*len(self.M)) # dates de fins des temps de blockage
            self.TBF = np.array([[0.0]*len(self.M[0])]*len(self.M)) # dates des debuts des temps de blockage
            # extraction des donnée depuis la matrice D
            char+='Calcul des temps de blocage\n'
            for i in range(len(self.D)-1):
                for j in range(len(self.D[0])):
                    char+='  TB['+str(i)+'][J'+str(self.S[j])+'] = D['+str(i+1)+'][J'+str(self.S[j])+'] - D['+str(i)+'][J'+str(self.S[j])+'] - P['+str(i)+'][J'+str(self.S[j])+'] = '+str(self.D[i+1][j])+' - '+str(self.D[i][j])+' - '+str(self.M[i][self.S[j]])+' = '+str(self.D[i+1][j]-self.D[i][j]-self.M[i][self.S[j]])+'\n'
                    self.F[i][j]=self.D[i][j]
                    self.DN[i][j]=self.D[i+1][j]
                    self.TB[i][j]=self.DN[i][j]-self.F[i][j]-self.M[i][self.S[j]]
                    self.O[i][j]=self.DN[i][j]-self.TB[i][j]
                    if self.TB[i][j]!=0:
                        self.TBO[i][j]=self.DN[i][j]
                        self.TBF[i][j]=self.O[i][j]
            if self.TT == 1:
                char2=""
                char3="Tardiness T:\n"
                char+="\n"+'Total flow time TFT = '
                TFT,Tm,TT=0.0,[],0.0
                #TFT : total flow time
                #Tm :temps de retard d'un job "tardiness"
                #TT : total tardiness
                for j in range(len(self.M[0])):
                    char+= '  D['+str(len(self.M))+'][J'+str(self.S[j])+']'
                    char2+= str(self.O[-1][j])
                    char3+= '  T[J'+str(self.S[j])+'] = '+'max( D['+str(len(self.M))+'][J'+str(self.S[j])+'] - d[J'+str(self.S[j])+'] , 0 ) = '+'max( '+str(self.D[-1][j])+' - '+str(self.dj[self.S[j]])+' , 0 ) = '+str(max(self.D[-1][j]-self.dj[self.S[j]],0))+'\n'
                    if j!=len(self.M[0])-1:
                        char+=" + "
                        char2+=" + "
                    else:
                        char+=" = "
                        char2+=" = "
                    TFT+=self.D[-1][j]
                    Tm.append(max(self.D[-1][j]-self.dj[self.S[j]],0))
                char+="\n"+char2
                char+=str(TFT)+'\n'
                char3+= "Tardiness TT = ∑ T[S[j]] : j = 0,...,"+str(len(self.M[0])-1)+"\n  TT = "+str(sum(Tm))
                char+="\n"+char3
                TT=sum(Tm)
                # renvoi des dates des fins et des debuts de chaque job et temp de blockage
                if self.do_report==0:
                    return [self.O,self.F,self.TBO,self.TBF,self.TB,TFT,Tm,TT]
                else:
                    return [self.O,self.F,self.TBO,self.TBF,self.TB,TFT,Tm,TT,self.D,char]
            else:
                char2=""
                char+="\n"+'Total flow time TFT = '
                TFT=0.0
                for j in range(len(self.M[0])):
                    char+= '  C['+str(len(self.M)-1)+'][J'+str(self.S[j])+']'
                    char2+= str(self.O[-1][j])
                    if j!=len(self.M[0])-1:
                        char+=" + "
                        char2+=" + "
                    else:
                        char+=" = "
                        char2+=" = "
                    TFT+=self.O[-1][j]
                char+="\n"+char2
                char+=str(TFT)
                if self.do_report==0:
                    return [self.O,self.F,self.TBO,self.TBF,self.TB,TFT]
                else:
                    return [self.O,self.F,self.TBO,self.TBF,self.TB,TFT,self.D,char]
        # avec preparaion et avec blockage
        elif self.prep == 1 and self.block == 1:
            return block_prep(self.M,list(self.S).copy(),self.Mprep,self.dj,self.TT,self.do_report) 
    # Méthode pour calculer les temps réels de fonctionnement TFR et d'arret TAR
    def TFR_TAR(self):
        TFR , TAR = [] , []
        #extraction de makspane
        Cmax = self.Cmax(self.S)
        # calcul de TFR/TAR pour le cas 1 -> preparation = arret
        for i in range(len(self.M)):
            tfr = 0 
            for j in range(len(self.M[0])):
                tfr += self.M[i][self.S[j]]
            cst = tfr/Cmax
            TFR.append(cst)
            TAR.append(1- cst)
        # calcul de TFR/TAR/TAP pour le cas 2 si on a de ppreparation -> preparation = travail
        if self.prep == 1:
            TFR_2 , TAR_2 , TAP = [] , [] , []
            for i in range(len(self.M)):
                tfr , tap = 0 , 0 
                for j in range(len(self.M[0])):
                    if j != 0:
                        tfr += self.M[i][self.S[j]]+self.Mprep[i][self.S[j-1]][self.S[j]]
                        tap += self.Mprep[i][self.S[j-1]][self.S[j]]
                    else:
                        tfr += self.M[i][self.S[j]]+self.Mprep[i][self.S[j]][self.S[j]]
                        tap += self.Mprep[i][self.S[j]][self.S[j]]
                cst = tfr/Cmax
                cst2 = tap/Cmax
                TFR_2.append(cst)
                TAR_2.append(1- cst)
                TAP.append(cst2)
            # si cas 2 renvoi tous les cas
            return [TFR,TAR,TFR_2,TAR_2,TAP]
        #si non renvoi le cas1
        return [TFR,TAR]
    # Méthode pour calculer les temps d'attente "utilusé dans les cas (0,1,0,0)/(1,1,0,0) "
    def TW(self):
        # TW :  temps d'attendre(time of waiting)
        # TWA :  temps d'attendre en arret (time of waiting arret)
        # TWA :  temps d'attendre en travail (time of waiting non arret)
        TW , TWA , TWNA = [] , [] , [[0.0]]
        traite = self.traite()  # donnée necessaire apres le traitement 
        # équation 0
        for i in range(len(self.M)-1):
            TW.append([])
            for j in range(len(self.M[0])):
                TW[i].append(traite[0][i+1][j]-self.M[i+1][self.S[j]]-traite[0][i][j])
        # équation 1
        for j in range(1,len(self.M[0])):
            TWNA[0].append(max(self.O[1][j-1]-self.O[0][j],0.0))
        # équation 2
        for i in range(1,len(self.M)-1):
            TWNA.append([])
            TWNA[i].append(self.F[i+1][0]-self.O[i][0])
        # équation 3
        for i in range(1,len(self.M)-1):
            TWNA.append([])
            for j in range(1,len(self.M[0])):
                TWNA[i].append(max(0.0,self.O[i+1][j-1]-self.O[i][j]))
        # équation 4
        TWNA.remove(TWNA[len(self.M)-1])
        for i in range(len(self.M)-1):
            TWA.append([])
            for j in range(len(self.M[0])):
                TWA[i].append(TW[i][j]-TWNA[i][j])
        return [TW,TWA,TWNA]
    # Méthode pour extracter le data frame pour tracer le Ganttchart
    def Gantt (self):
        global color_names       # liste des couleur utilusée sans préparaion ou avec noidel ou avec nowait
        global color_names_prep  # liste des couleur utilusée avec préparaion
        Matrix = self.traite()   # données traitées
        Cmax= Matrix[0][len(Matrix[0])-1][-1]
        # s'il y'a une erreur dans le traitement
        if type(Matrix)==str:
            return Matrix
        START , END , Machines , colors , texts= [] , [] , [] , [] , [] 
        # START : list des temps de débuts de chaque block soit temps de processus ou de preparaion ou de blocage
        # END : list des temps de fin de chaque block soit temps de processus ou de preparaion ou de blocage
        # Machines : list des labels de chaque machine
        # colors : list des couleur de chaque block soit temps de processus ou de preparaion ou de blocage
            # noir pour les temps de preparation / gris pour blocage / autres pour les temps de processus
        # text : labels de chaque block
            # Jj pour les jobs / TP pour les préparaion / TB pour les blocages
        Gantt = {'Machines' : None , 'Start' : None,'End' : None , 'Color' : None, 'Label': None }
        # Gantt : dictionnaire des données utilusée pour afficher le gaint
        for i in range(len(self.M)-1,-1,-1):  # remplissage des labels pour les machines 
            Machines.append("Machine "+str(i))
        # sans preparation et sans blockage ou nowait ou noidel
        if (self.prep == 0 and self.block == 0) or self.noidel==1 or self.nowait == 1:
            if len(Matrix[0][0])> len(color_names):
                color_names = color_names*((len(Matrix[0][0])//len(color_names))+1)
            for lst1,lst2,k in zip(Matrix[0],Matrix[1],range(len(self.M))):
                START.append([])
                END.append([])
                colors.append([])
                texts.append([])
                for e1,e2,c,i in zip(lst1,lst2,color_names,range(len(lst1))):
                    START[k].append([e2])
                    END[k].append([e1])
                    colors[k].append(c)
                    texts[k].append("J"+str(self.S[i]))
        # avec preparaion et sans blockage
        elif self.prep == 1 and self.block == 0:
            M1 , M2 = [] , []
            for lst1,lst2,lst3,lst4 in zip(Matrix[0],Matrix[1],Matrix[2],Matrix[3]):
                lst5 , lst6 = lst1+lst3 , lst2+lst4
                lst5.sort()
                lst6.sort()
                M1.append(lst5)
                M2.append(lst6)
            if len(M1[0])> len(color_names_prep):
                color_names_prep = color_names_prep*((len(Matrix[0][0])//len(color_names_prep))+1)
            for lst1,lst2,k in zip(M1,M2,range(len(M1))):
                START.append([])
                END.append([])
                colors.append([])
                texts.append([])
                for e1,e2,c,i in zip(lst1,lst2,color_names_prep,range(len(lst1))):
                    START[k].append([e2])
                    END[k].append([e1])
                    colors[k].append(c)
                    if i%2 == 1:
                        texts[k].append("J"+str(self.S[i//2]))
                    else:
                        texts[k].append("TP")
        # sans preparation et avec blockage 
        elif self.prep == 0 and self.block == 1:
            if len(Matrix[0][0])> len(color_names):
                color_names = color_names*((len(Matrix[0][0])//len(color_names))+1)
            for Jends,Jstarts,TBends,TBstarts,TBs,i in zip(Matrix[0],Matrix[1],Matrix[2],Matrix[3],Matrix[4],range(len(Matrix[0]))):
                START.append([])
                END.append([])
                colors.append([])
                texts.append([])
                for Jend,Jstart,TBend,TBstart in zip(Jends,Jstarts,TBends,TBstarts):
                    START[i].append([float(Jstart)])
                    END[i].append([float(Jend)])
                    if TBstart!=0:
                        START[i].append([float(TBstart)])
                        END[i].append([float(TBend)])
                START[i].sort()
                END[i].sort()
                index_color = 0
                colors[i].append(color_names[index_color])
                texts[i].append("J"+str(self.S[index_color]))
                index_color+=1
                for j in range(1,len(START[i])):
                    if START[i][j] in TBstarts:
                        colors[i].append('silver')
                        texts[i].append("TB")
                    else:
                        colors[i].append(color_names[index_color])
                        texts[i].append("J"+str(self.S[index_color]))
                        index_color+=1
        # avec preparaion et avec blockage            
        elif self.prep == 1 and self.block == 1:
            if len(Matrix[0][0])> len(color_names):
                color_names = color_names*((len(Matrix[0][0])//len(color_names))+1)
            for Jends,Jstarts,Pends,Pstarts,TBends,TBstarts,i in zip(Matrix[0],Matrix[1],Matrix[2],Matrix[3],Matrix[4],Matrix[5],range(len(Matrix[0]))):
                START.append([])
                END.append([])
                colors.append([])
                texts.append([])
                for Jend,Jstart,Pend,Pstart,TBend,TBstart,j in zip(Jends,Jstarts,Pends,Pstarts,TBends,TBstarts,range(len(Matrix[0][0]))):
                    START[i].append([float(Pstart)])
                    texts[i].append("TP")
                    colors[i].append("black")
                    START[i].append([float(Jstart)])
                    texts[i].append("J"+str(self.S[j]))
                    colors[i].append(color_names[j])
                    END[i].append([float(Pend)])
                    END[i].append([float(Jend)])
                    if TBstart!=0:
                        print("added")
                        START[i].append([float(TBstart)])
                        texts[i].append("TB")
                        colors[i].append('silver')
                        END[i].append([float(TBend)])
        START.reverse()
        END.reverse()
        texts.reverse()
        colors.reverse()
        Gantt['Machines'] = Machines
        Gantt['Start'] = START
        Gantt['End'] = END
        Gantt['Color'] = colors
        Gantt['Label'] = texts
        tasks_df = pd.DataFrame(Gantt)
        tasks_df = tasks_df.set_index('Machines')
        if self.sim==0:
            return tasks_df
        else:
            return [Gantt,self.M,self.S,Cmax]
    # Méthode pour extracter le Cmax
    def Cmax(self,Seq):
        self.S = Seq
        return max(self.traite()[0][len(self.M)-1])
    # Méthode pour extracter le Cmax et TT(total tardiness)
    def Cmax_TT(self,Seq):
        self.S = Seq
        donees = self.traite()    
        return [max(donees[0][len(self.M)-1]),donees[-1]] # retourner Cmax et TT
    # Méthode pour calculer les combinaisons possibles
    def combinations (self) :
        n = len(self.M[0])
        combinations = generate_combinations(n)
        Cmaxs = []
        if self.TT ==1:
            TTs = []
            for comb in combinations:
                cmax_tt = self.Cmax_TT(comb)
                Cmaxs.append(cmax_tt[0])
                TTs.append(cmax_tt[1])
            return [Cmaxs,combinations,TTs]
        else:
            for comb in combinations:
                Cmaxs.append(self.Cmax(comb))
            return [Cmaxs,combinations]
    # Méthode pour calculer les combinaisons où Cmax = min(Cmaxs)
    def choosed_combinations(self):
        lst = self.combinations()
        Cmaxs = np.array(lst[0])
        indexs = np.where(Cmaxs == min(Cmaxs))
        choosed = []
        for i in indexs[0]:
            choosed.append(lst[1][i])
        return [min(Cmaxs),choosed]
    # Méthode pour calculer les combinaisons où TT = min(TTs)
    def choosed_combinations_by_TT(self):
        lst = self.combinations()
        TTs = np.array(lst[2])
        indexs = np.where( TTs == min(TTs))
        choosed = []
        for i in indexs[0]:
            choosed.append(lst[1][i])
        return [min(TTs),choosed]
    # Méthode pour calculer les combinaisons où TT = min(TTs) et Cmax = min(Cmaxs)
    def choosed_combinations_by_TT_Cmax(self):
        lst = self.combinations()
        TTs = np.array(lst[2])
        Cmaxs = np.array(lst[0])
        indexs_tt = np.where( TTs == min(TTs))
        indexs_cmax = np.where(Cmaxs == min(Cmaxs))
        choosed_by_cmax = []
        choosed_by_tt = []
        for i in indexs_cmax[0]:
            choosed_by_cmax.append(lst[1][i])
        for i in indexs_tt[0]:
            choosed_by_tt.append(lst[1][i])
        choosed = []
        for comb in choosed_by_tt:
            if comb in choosed_by_tt:
                choosed.append(comb)
        if choosed!=[]:
            return [(min(TTs),min(Cmaxs)),choosed]
        else:
            return "il n'y a pas de séquence vérifiant min(TT) et min(Cmax)"
    # Méthode pour calculer la sequence optimale dans le cas de deux machines
    def Johnson(self): #(Algorithme de jonhson)
        if len(self.M) !=2 :
            print("Jonhson Algorithme ne fonctionne qu'avec deux machines")
            return None
        # etape 1 et 2
        jobs = list(range(len(self.M[0])))
        U,V = [],[]
        for j,t1,t2 in zip(jobs,self.M[0],self.M[1]):
            if t1 < t2:
                U.append(j)
            else:
                V.append(j)
        X,Y = U.copy(),V.copy()
        # etape 3
        SPT,LPT,M1,M2 = [],[],[],[]
        for j in U:
            M1.append(self.M[0][j])
        for j in V:
            M2.append(self.M[1][j])
        x = 0
        n= len (U)
        MA,MB = [],[]
        while 1:
            if len(SPT)==n:
                break
            if M1[x] == min(M1):
                SPT.append(U[x])
                MA.append(M1[x])
                M1.remove(M1[x])
                U.remove(U[x])
                x=0
            else:
                x+=1
        x = 0
        n= len (V)
        while 1:
            if len(LPT)==n:
                break
            if M2[x] == max(M2):
                LPT.append(V[x])
                MB.append(M2[x])
                M2.remove(M2[x])
                V.remove(V[x])
                x=0
            else:
                x+=1
        print(SPT,LPT)
        # etape 4
        A1=find_combinations(SPT, MA)
        A2=find_combinations(LPT, MB)
        RESULT =combine_lists(A1,A2)
        return [X,Y,SPT,LPT,RESULT]
    # Méthode pour calculer la sequence optimale dans le cas de plusieurs machines 
    def CDS(self,k): # (CAMPBELL, DUDEK et SMITH)
        if len(self.M) <= 2 :
            print("CDS ne fonctionne qu'avec plus de deux machines")
            return None
        if len(self.M) == 2 :
            print("on va applique Jonhson car il y'a deux machines ")
            return self.Johnson()
        Matrix = [[],[]]
        for j in range(len(self.M[0])):
            som = 0
            for i in range(k):
                som = som + self.M[i][j]
            Matrix[0].append(som)
        for j in range(len(self.M[0])):
            som = 0
            for i in range(len(self.M)-k,len(self.M)):
                som = som + self.M[i][j]
            Matrix[1].append(som)
        self.M = Matrix
        if self.do_report==0:
            return self.Johnson()
        else:
            return [self.Johnson(),Matrix]
# fonction qui traite les problemes flowshop avec condition de blocage avec les temps de préparation
def block_prep(P,S,TP,dj,T,report):
    #preparation des variables et equation 1 et 2 "adéquation des systèmes d'indexation"
    char = "Avec préparation et avec blocage \n on Ajout une machine vertuel et un job virtuel pour calculer la matrice D\n indéxation des machines réeel de 1 à "+str(len(P))+' ; et des jobs de 1 à '+str(len(P[0]))+'\n'
    char+='équation 1: pour i = 0,...,'+str(len(P))+' et j = 0\n'
    char+='  D[i][J0] = 0\n'
    for i in range(len(S)):
        S[i]=S[i]+1
    S = np.array(S) 
    S = np.hstack(([0], S))
    new_line = [0.0]*len(P[0])
    P = np.vstack((new_line, P))
    new_column = [[0.0]]*len(P)
    P = np.hstack((new_column,P)) 
    D = np.array([[0.0]*(len(P[0]+1))]*(len(P)))
    #equation 2 3 et 4
    for j in range(1,len(D[0])):
        for i in range(len(D)):
            if i == 0:
                # équation 2 et 3
                if j == 1:
                   char+='équation 2: pour i = 0 et j = 1\n'
                   char+='  D[0][J'+str(S[j])+'] = D[1][J0] + M[1][J'+str(S[1])+'][J'+str(S[1])+'] = '+str(D[1][j-1])+' + '+str(TP[0][S[1]-1][S[1]-1] )+' = '+str(D[1][j-1] + TP[0][S[1]-1][S[1]-1] )+'\n'
                   D[0][j] = D[1][j-1] + TP[0][S[1]-1][S[1]-1] 
                else:
                    char+='équation 3: pour i = 0 et j = '+str(j)+'\n'
                    char+='  D[0][J'+str(S[j])+'] = D[1][J'+str(S[j-1])+'] + M[1][J'+str(S[j-1])+'][J'+str(S[j])+'] = '+str(D[1][j-1])+' + '+str(TP[0][S[j-1]-1][S[j]-1])+' = '+str(D[1][j-1] + TP[0][S[j-1]-1][S[j]-1])+'\n'
                    D[0][j] = D[1][j-1] + TP[0][S[j-1]-1][S[j]-1]
            elif i == len(D)-1:
                # équation 6
                char+='équation 6: pour i = '+str(len(D)-1)+' et j = '+str(j)+'\n'
                char+='  D['+str(len(D)-1)+'][J'+str(S[j])+'] = D['+str(len(D)-2)+'][J'+str(S[j])+'] + P['+str(len(P)-1)+'][J'+str(S[j])+'] = '+str(D[-2][j])+' + '+str(P[-1][S[j]])+' = '+str(D[-2][j] + P[-1][S[j]])+'\n'
                D[-1][j] = D[-2][j] + P[-1][S[j]]
            else:
                # équation 4
                if j == 1:
                    char+='équation 4: pour i = '+str(i)+' et j = 1\n'
                    char+='  D['+str(i)+'][J'+str(S[j])+'] = max( D['+str(i-1)+'][J'+str(S[j])+'] + P['+str(i)+'][J'+str(S[j])+'] , D['+str(i+1)+'][J0] + M['+str(i)+'][J'+str(S[j])+'][J'+str(S[j])+'] ) = max( '+str(D[i-1][1])+' + '+str(P[i][S[1]])+' , ' + str(D[i+1][0])+' + '+str(TP[i][S[1]-1][S[1]-1])+' ) = '+str(max(D[i-1][1] + P[i][S[1]] , D[i+1][0] + TP[i][S[1]-1][S[1]-1]))+'\n'
                    D[i][1] = max(D[i-1][1] + P[i][S[1]] , D[i+1][0] + TP[i][S[1]-1][S[1]-1])
                # équation 5
                else:
                    char+='équation 5: pour i = '+str(i)+' et j = '+str(j)+'\n'
                    char+='  D['+str(i)+'][J'+str(S[j])+'] = max( D['+str(i-1)+'][J'+str(S[j])+'] + P['+str(i)+'][J'+str(S[j])+'] , D['+str(i+1)+']['+str(S[j-1])+'] + M['+str(i)+']['+str(S[j-1])+']['+str(S[j])+'] ) = max( '+str(D[i-1][j])+' + '+str(P[i][S[j]])+' , ' + str(D[i+1][j-1])+' + '+str(TP[i][S[j-1]-1][S[j]-1])+' ) = '+str(max(D[i-1][j] + P[i][S[j]] , D[i+1][j-1] + TP[i][S[j-1]-1][S[j]-1]))
                    D[i][j] = max(D[i-1][j] + P[i][S[j]] , D[i+1][j-1] + TP[i][S[j-1]-1][S[j]-1])
    O,F,OP,FP,OB,FB,OOB,TB = [],[],[],[],[],[],[],[]
    # O: dates des fins des jobs (out of job)
    # F: dates des debuts des jobs (first of job)
    # OP : dates des fins des temps de préparation (out of preparation time)
    # FP : dates des debuts des temps de préparation (firt of preparation time)
    # FB : dates des debuts des temps de blocage (first of blocing time)
    # OB : dates des find des temps de blocage (first of job blocking time)
    # OOB : partie supérieur à droite de D
    # TB : temps de blocage (blocking time) 
    #préparation pour le calcul
    for i in range(len(D)):
        Flst,FPlst,OOBlst=[],[],[]
        for j in range(len(D[0])):
            if i != 0 and j != len(D[0])-1:
                FPlst.append(D[i][j])
            if i != len(D)-1 and j != 0:
                Flst.append(D[i][j])
            if i !=0 and j !=0:
                OOBlst.append(D[i][j])
        if i != 0 :
            FP.append(FPlst)
            OOB.append(OOBlst)
        if i != len(D)-1:
            F.append(Flst)
    # calcul des matrices nécessaire pour tracer le gaint
    char+= ' Indexation normale pour TB et C\n'
    charTB='Calcul des temps de blocage\n'
    charC='Calcul de la matrice C\n'
    for i in range(len(F)):
        O.append([])
        OP.append([])
        TB.append([])
        OB.append([])
        FB.append([])
        for j in range(len(F[0])): 
            O[i].append(F[i][j]+P[i+1][S[j+1]])
            charC+='  C['+str(i)+'][J'+str(S[j+1]-1)+'] = D['+str(i+1)+'][J'+str(S[j])+'] + P['+str(i)+'][J'+str(S[j+1]-1)+'] = '+str(F[i][j])+' + '+str(P[i+1][S[j+1]])+' = '+str(F[i][j]+P[i+1][S[j+1]])+'\n'
            if j == 0:
                OP[i].append(FP[i][0]+TP[i][S[1]-1][S[1]-1])
            else:
                OP[i].append(FP[i][j]+TP[i][S[j]-1][S[j+1]-1])
            TB[i].append(OOB[i][j]-O[i][j])
            charTB+='  TB['+str(i)+'][J'+str(S[j+1]-1)+'] = D['+str(i+1)+'][J'+str(S[j])+'] - C['+str(i)+'][J'+str(S[j+1]-1)+'] = '+str(OOB[i][j])+' - '+str(O[i][j])+' = '+str(OOB[i][j]-O[i][j])+'\n'
            if TB[i][j] == 0:
                OB[i].append(0.0)
                FB[i].append(0.0)
            else:
                FB[i].append(O[i][j])
                if j != len(F[0])-1:
                    OB[i].append(FP[i][j+1]) 
                else:
                    OB[i].append(O[i][j]+TB[i][j]) 
    char+=charC
    char+=charTB
    TBP,TBB,TBF = [],[],[]
    # TBB : temps de bolocage à cause d'un autre blocage 
    # TBP : temps de bolocage à cause de préparation
    # TBF : temps de blocage à cause de travail
    # calcul des temps de blocage
    for i in range(len(O)-1):
        TBP.append([])
        TBB.append([0.0])
        TBF.append([0.0])
        if TB[i][0] != 0:
            TBP[i].append(TB[i][0]) 
        else:
            TBP[i].append(0.0) 
    for i in range(len(O)-1):
        for j in range(1,len(O[0])):
            if TB[i][j] != 0:
                if F[i+1][j]-FP[i+1][j] >= TB[i][j]:
                    TBP[i].append(TB[i][j])
                    TBF[i].append(0.0)
                    TBB[i].append(0.0)
                else:
                    TBP[i].append(TP[i+1][S[j]-1][S[j+1]-1])
                    TBF[i].append(O[i+1][j-1] - O[i][j])
                    if TB[i+1][j-1] != 0:
                        TBB[i].append(TB[i+1][j-1])
                    else:
                        TBB[i].append(0.0)
            else:
                TBP[i].append(0.0)
                TBB[i].append(0.0)
                TBF[i].append(0.0)
    if T==1:
        char2=""
        char3="Tardiness T:\n"
        char+="\n"+'Total flow time TFT = '
        TFT,Tm,TT=0.0,[],0.0
        #TFT : total flow time
        #Tm :temps de retard d'un job "tardiness"
        #TT : total tardiness
        for j in range(len(P[0])):
            char+= 'D['+str(len(P))+'][J'+str(S[j])+']'
            char2+= str(D[-1][j])
            char3+= '  T[J'+str(S[j])+'] = '+'max( D['+str(len(P))+'][J'+str(S[j])+'] - d[J'+str(S[j]-1)+'] , 0 ) = '+'max( '+str(D[-1][j])+' - '+str(dj[S[j]-1])+' , 0 ) = '+str(max(D[-1][j]-dj[S[j]-1],0))+'\n'
            if j!=len(P[0])-1:
                char+=" + "
                char2+=" + "
            else:
                char+=" = "
                char2+=" = "
            TFT+=D[-1][j]
            if j!=0:
                Tm.append(max(D[-1][j]-dj[S[j]-1],0))
        char+="\n"+char2
        char+=str(TFT)+'\n'
        char3+= "Tardiness TT = ∑ T[S[j]] : j = 0,...,"+str(len(P[0])-1)+"\n  TT = "+str(sum(Tm))
        char+="\n"+char3
        TT=sum(Tm)
        if report==0:
            return [O,F,OP,FP,OB,FB,TB,TBF,TBB,TBP,TFT,Tm,TT]
        else:
            return [O,F,OP,FP,OB,FB,TB,TBF,TBB,TBP,TFT,Tm,TT,D,char]
    else:
        char2=""
        char+="\n"+'Total flow time TFT = '
        TFT=0.0
        for j in range(len(P[0])):
            char+= '  C['+str(len(P)-1)+'][J'+str(S[j]-1)+']'
            char2+= str(O[-1][j])
            if j!=len(P[0])-1:
                char+=" + "
                char2+=" + "
            else:
                char+=" = "
                char2+=" = "
            TFT+=D[-1][j]
        char+="\n"+char2
        char+=str(TFT)
        if report==0:
            return [O,F,OP,FP,OB,FB,TB,TBF,TBB,TBP,TFT]
        else:
            return [O,F,OP,FP,OB,FB,TB,TBF,TBB,TBP,TFT,D,char]
# fonction qui traite les problemes flowshop avec condition de non arret 
def noidel_probleme_solving(P,S,dj,T,report):
    L,C,F,TW=[0.0],[],[],[] 
    char = 'Sans préparation et dans arrets\n'
    # L list des temps mort dans chaque machine et équation 1
    # C matrice des temps de fin de chaque job
    # F matrice des temps des début de chaque job
    # TW matrice des temps d'attete de chaque job entre machine i et i+1 
    n,m=len(P[0]),len(P) 
    # n est le nombre total des jobs et m et le nombre total des machines
    # application des formules mathématique
    char+='équation 1 : pour i = 0 : \n L[0] = 0\n'
    for i in range(m): 
        C.append([])
        F.append([])
        # équation 2
        if i!=0:
            char+='équation 2: pour i='+str(i)+'\n'
            charmaxs="max( "
            charmaxsnum="max( "
            maxs = []
            for k in range(n):
                char+='  maxs['+str(k)+'] =  '
                mx1,mx2 = 0,0
                sum1,sum11,sum2,sum22='','','',''
                for j in range(n-k):
                    sum1+='P['+str(i-1)+'][J'+str(S[j])+']'
                    sum11+=str(P[i-1][S[j]])
                    mx1+=P[i-1][S[j]]
                    if j != n-k-1:
                        sum1+= ' + '
                        sum11+= ' + ' 
                for j in range(n-k-1):
                    sum2+='P['+str(i)+'][J'+str(S[j])+']'
                    sum22+=str(P[i][S[j]])
                    mx2+=P[i][S[j]]
                    if j != n-k-2:
                        sum2+= ' - '
                        sum22+= ' - ' 
                char+=sum1+' - '+sum2+ ' = '+sum11+' - '+sum22+' = '+str(mx1-mx2)+'\n'
                charmaxs+='maxs['+str(k)+']'
                charmaxsnum+=str(mx1-mx2)
                if k != n-1:
                    charmaxs+= ' , '
                    charmaxsnum+= ' , '
                maxs.append(mx1-mx2)
            char+='  L['+str(i)+'] = L['+str(i-1)+'] + '+charmaxs+' ) = '+str(L[i-1])+' + '+charmaxsnum+' ) = '+str(L[i-1])+' + '+str(max(maxs))+' = '+str(L[i-1]+max(maxs))+'\n'
            L.append(L[i-1]+max(maxs))
        char+='équation 3 et 4 : pour i = '+str(i)+' et j = 0,...,'+str(n-1)+': \n'
        for j in range(n):
            # équation 3
            if j==0:
                char+='  C['+str(i)+'][J'+str(S[j])+'] = L['+str(i)+'] + P['+str(i)+'][J'+str(S[j])+'] = '+str(L[i])+' + '+str(P[i][S[j]])+' = '+str(L[i]+P[i][S[j]])+'\n'
                C[i].append(L[i]+P[i][S[j]])
                F[i].append(L[i])
            # équation 4
            else:
                char+='  C['+str(i)+'][J'+str(S[j])+'] = C['+str(i)+'][J'+str(S[j-1])+'] + P['+str(i)+'][J'+str(S[j])+'] = '+str(C[i][j-1])+' + '+str(P[i][S[j]])+' = '+str(C[i][j-1]+P[i][S[j]])+'\n'
                C[i].append(C[i][j-1]+P[i][S[j]])
                F[i].append(C[i][j-1])
    # calcul des temps d'attendre de chaque job de la machine i à i+1
    char+="Calcul des temps d'attendre :\n"
    for i in range(m-1):
        TW.append([])
        for j in range(n):
            char+='  TA['+str(i)+'/'+str(i+1)+'][J'+str(S[j])+'] = C['+str(i+1)+'][J'+str(S[j])+'] - P['+str(i+1)+'][J'+str(S[j])+'] - C[i][J'+str(S[j])+'] = '+str(C[i+1][j])+' - '+str(P[i+1][j])+' - '+str(C[i][j])+' = '+str(F[i+1][j]-C[i][j])+'\n'
            TW[i].append(F[i+1][j]-C[i][j])
    if T==1:
        char2=""
        char3="Tardiness T:\n"
        char+="\n"+'Total flow time TFT = '
        TFT,Tm,TT=0.0,[],0.0
        #TFT : total flow time
        #Tm :temps de retard d'un job "tardiness"
        #TT : total tardiness
        for j in range(n):
            char+= '  C['+str(m-1)+'][J'+str(S[j])+']'
            char2+= str(C[-1][j])
            char3+= '  T[J'+str(S[j])+'] = '+'max( C['+str(m-1)+'][J'+str(S[j])+'] - d[J'+str(S[j])+'] , 0 ) = '+'max( '+str(C[-1][j])+' - '+str(dj[S[j]])+' , 0 ) = '+str(max(C[-1][j]-dj[S[j]],0))+'\n'
            if j!=n-1:
                char+=" + "
                char2+=" + "
            else:
                char+=" = "
                char2+=" = "
            TFT+=C[-1][j]
            Tm.append(max(C[-1][j]-dj[S[j]],0))
        char+="\n"+char2
        char+=str(TFT)+'\n'
        char3+= "Tardiness TT = ∑ T[S[j]] : j = 0,...,"+str(n-1)+"\n  TT = "+str(sum(Tm))
        char+="\n"+char3
        TT=sum(Tm)
        if report==0:
            return [C,F,TW,TFT,Tm,TT]
        else:
            return [C,F,TW,TFT,Tm,TT,char]
    else:
        char2=""
        char+="\n"+'Total flow time TFT = '
        TFT=0.0
        for j in range(n):
            char+= '  C['+str(m-1)+'][J'+str(S[j])+']'
            char2+= str(C[-1][j])
            if j!=n-1:
                char+=" + "
                char2+=" + "
            else:
                char+=" = "
                char2+=" = "
            TFT+=C[-1][j]
        char+="\n"+char2
        char+=str(TFT)
        if report==0:
            return [C,F,TW,TFT]
        else:
            return [C,F,TW,TFT,char]
# fonction qui traite les problemes flowshop avec condition de non attendre
def nowait_probleme_solving(P,S,dj,T,report):
    char = 'Sans préparation et dans attendre\n'
    n,m=len(P[0]),len(P) # n est le nombre total des jobs et m et le nombre total des machines
    D = np.array([[0.0]*n]*n) # matrice des delais entre deux jobs j1 et j2
    C = [] # C matrice des temps de fin de chaque job "au début elle est seulement la derniere ligne"
    # équation 1 : calcul de la matrice des delais
    char+='équation 1 : Calcul de la matrice des délais D entre chaque deux job tel que j1 = 0,...,'+str(n-1)+' j2 = 0,...,'+str(n-1)+' :\n'
    for j1 in range(n):
        for j2 in range(n):
            mxs = []
            for r in range(m):
                char+='  maxs['+str(r)+'] = '
                mx1,mx2=0,0
                mx1char,mx2char='',''
                mx1charnum,mx2charnum='',''
                for k in range(1,r+1):
                    mx1char+='P['+str(k)+'][J'+str(j1)+']'
                    mx1charnum+=str(P[k][j1])
                    if k != r:
                        mx1char+=' + '
                        mx1charnum+=' + '
                    mx1+=P[k][j1]
                for k in range(r):
                    mx2char+='P['+str(k)+'][J'+str(j2)+']'
                    mx2charnum+=str(P[k][j2])
                    if k != r-1:
                        mx2char+=' - '
                        mx2charnum+=' - '
                    mx2+=P[k][j2]
                if mx1char=='':
                    mx1char='0'
                if mx2char=='':
                    mx2char='0'
                if mx1charnum=='':
                    mx1charnum='0'
                if mx2charnum=='':
                    mx2charnum='0'
                if mx1char==mx1charnum and mx2char==mx2charnum:
                    char+=mx1charnum+' - '+mx2charnum+" = "+str(mx1-mx2)+'\n'
                elif str(mx1)==mx1charnum and str(mx2)==mx2charnum :
                    char+=mx1char+' - '+mx2char+" = "+mx1charnum+' - '+mx2charnum+" = "+str(mx1-mx2)+'\n'
                else:
                    char+=mx1char+' - '+mx2char+" = "+mx1charnum+' - '+mx2charnum+" = "+str(mx1)+' - '+str(mx2)+' = '+str(mx1-mx2)+'\n'
                mxs.append(mx1-mx2)
            char+='  maxs = '+str(mxs)+'\n'
            char+='  D[J'+str(j1)+'][J'+str(j2)+'] = P[0][J'+str(j1)+'] + max(max(maxs),0)'+' = '+str(P[0][j1])+' + '+str(max(max(mxs),0))+' = '+str(P[0][j1] + max(max(mxs),0))+'\n'
            D[j1][j2] = P[0][j1] + max(max(mxs),0)
    # équation 2 : calcul des temps de fins et de début de chaque job dans la dernière machine
    char+='équation 2 : Calcul de la dernière ligne de la matrice C pour i = '+str(m-1)+' et j = 0,...,'+str(n-1)+' :\n'
    for j in range(n):
        v1,v2=0,0
        v1char,v2char="",""
        v1charnum,v2charnum="",""
        for k in range(1,j+1):
            v1char+='D[J'+str(S[k-1])+'][J'+str(S[k])+']'
            v1charnum+=str(D[S[k-1]][S[k]])
            if k!=j:
                v1char+=' + '
                v1charnum+=' + '
            v1+=D[S[k-1]][S[k]]
        for k in range(m):
            v2char+='P['+str(k)+'][J'+str(S[j])+']'
            v2charnum+=str(P[k][S[j]])
            if k!=m-1:
                v2char+=' + '
                v2charnum+=' + '
            v2+=P[k][S[j]]
        if v1char=='':
            char+='C['+str(m-1)+'][J'+str(S[j])+'] = '+v2char+' = '+v2charnum+' = '+str(v2)+" = "+str(v1+v2)+'\n'
        elif v2char=='':
            char+='C['+str(m-1)+'][J'+str(S[j])+'] = '+v1char+' = '+v1charnum+' = '+str(v1)+" = "+str(v1+v2)+'\n'
        else:
            char+='C['+str(m-1)+'][J'+str(S[j])+'] = '+v1char+' + '+v2char+' = '+v1charnum+' + '+v2charnum+' = '+str(v1)+' + '+str(v2)+" = "+str(v1+v2)+'\n'
        C.append(v1+v2)
    C=np.array([[0.0]*n]*(m-1)+[C]) # C devien une matrice
    F=np.array([[0.0]*n]*(m))  # F matrice des temps des début de chaque job
    char+='équation 3 : Calcul de la matrice C pour i = '+str(m-2)+',...,0 et j = 0,...,'+str(n-1)+' :\n'
    for i in range(m-1,-1,-1):
        for j in range(n):
            if i!=m-1:
                char+='C['+str(i)+'][J'+str(S[j])+'] = C['+str(i+1)+'][J'+str(S[j])+'] - P['+str(i+1)+'][J'+str(S[j])+'] = '+str(C[i+1][j])+' - '+str(P[i+1][S[j]])+' = '+str(C[i+1][j]-P[i+1][S[j]])+'\n'
                C[i][j]=C[i+1][j]-P[i+1][S[j]]
            F[i][j] = C[i][j] - P[i][S[j]]
    TW=np.array([[0.0]*(n-1)]*(m)) 
    # TW : matrice des temps d'attente entre deux jobs successif dans chaque machine
    for i in range(m):
        for j in range(n-1):
            TW[i][j] = F[i][j+1]-C[i][j]
    if T ==1:
        char2=""
        char3="Tardiness T:\n"
        char+="\n"+'Total flow time TFT = '
        TFT,Tm,TT=0.0,[],0.0
        #TFT : total flow time
        #Tm :temps de retard d'un job "tardiness"
        #TT : total tardiness
        for j in range(n):
            char+= '  C['+str(m-1)+'][J'+str(S[j])+']'
            char2+= str(C[-1][j])
            char3+= '  T[J'+str(S[j])+'] = '+'max( C['+str(m-1)+'][J'+str(S[j])+'] - d[J'+str(S[j])+'] , 0 ) = '+'max( '+str(C[-1][j])+' - '+str(dj[S[j]])+' , 0 ) = '+str(max(C[-1][j]-dj[S[j]],0))+'\n'
            if j!=n-1:
                char+=" + "
                char2+=" + "
            else:
                char+=" = "
                char2+=" = "
            TFT+=C[-1][j]
            Tm.append(max(C[-1][j]-dj[S[j]],0))
        char+="\n"+char2
        char+=str(TFT)+'\n'
        char3+= "Tardiness TT = ∑ T[S[j]] : j = 0,...,"+str(n-1)+"\n  TT = "+str(sum(Tm))
        char+="\n"+char3
        TT=sum(Tm)
        if report ==0:
            return [C,F,TW,TFT,Tm,TT]
        else:
            return [C,F,TW,TFT,Tm,TT,char]
    else:
        char2=""
        char+="\n"+'Total flow time TFT = '
        TFT=0.0
        for j in range(n):
            char+= '  C['+str(m-1)+'][J'+str(S[j])+']'
            char2+= str(C[-1][j])
            if j!=n-1:
                char+=" + "
                char2+=" + "
            else:
                char+=" = "
                char2+=" = "
            TFT+=C[-1][j]
        char+="\n"+char2
        char+=str(TFT)

        if report==0:
            return [C,F,TW,TFT]
        else:
            return [C,F,TW,TFT,char]






































color_names = ['red', 'green', 'blue', 'navy' , 'orange', 'purple', 'pink', 'brown', 'cyan', 'magenta',
               'teal', 'lime', 'lavender', 'indigo', 'maroon', 'gold', 'olive', 'yellow']
color_names_prep = ['black','red','black', 'green','black', 'blue','black', 'yellow','black', 'orange','black', 'purple','black', 'pink','black', 'brown','black', 'cyan','black', 'magenta','black',
                    'teal','black', 'lime','black', 'lavender','black', 'indigo','black', 'maroon','black', 'gold','black', 'olive','black', 'navy','black', 'silver']

Q = [[2,4,5,3],[5,6,7,8],[3,2,3,2]]
H = [[9,3,5,6,4,8,6,2],[7,7,4,8,7,5,9,5]]
P = [[5,2,3,6,8,4,5,7],[2,4,4,5,3,9,6,3],[3,2,5,4,2,7,8,2]]
s = [2, 5, 6, 3, 0, 1, 4, 7]



class App(ct.CTk):
    def __init__(self):
        super().__init__()
        self.title("Ordonnancement")
        self.geometry("800x610")
        self.minsize(800, 610)
        self.maxsize(800, 610)
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)

        image_path = os.path.join(os.path.dirname(os.path.realpath(__file__)), "images")
        # create navigation frame
        self.navigation_frame = ct.CTkFrame(self, corner_radius=0)
        self.navigation_frame.grid(row=0, column=0, sticky="nsew")
        self.navigation_frame.grid_rowconfigure(4, weight=1)
         # load images with light and dark mode image
        self.logo_image = ct.CTkImage(light_image=Image.open(os.path.join(image_path, "logo_light.png")),
                                      dark_image=Image.open(os.path.join(image_path, "logo_dark.png")), size=(40, 40))
        self.combinations_image = ct.CTkImage(light_image=Image.open(os.path.join(image_path, "combinations_light.webp")),
                                              dark_image=Image.open(os.path.join(image_path, "combinations_dark.png")), size=(40, 40))
        self.matrix_image = ct.CTkImage(light_image=Image.open(os.path.join(image_path, "matrix_light.webp")),
                                        dark_image= Image.open(os.path.join(image_path, "matrix_dark.png")), size=(40, 40))
        self.home_image = ct.CTkImage(light_image=Image.open(os.path.join(image_path, "home_dark.png")),
                                                 dark_image=Image.open(os.path.join(image_path, "home_light.png")), size=(40, 40))
        self.Gantt_image = ct.CTkImage(light_image=Image.open(os.path.join(image_path, "Gantt_light.png")),
                                                 dark_image=Image.open(os.path.join(image_path, "Gantt_dark.png")), size=(40, 40))
        self.sequence_image = ct.CTkImage(light_image=Image.open(os.path.join(image_path, "sequence_light.webp")),
                                                     dark_image=Image.open(os.path.join(image_path, "sequence_dark.png")), size=(40, 40))
        self.in_home_image = ct.CTkImage(Image.open(os.path.join(image_path, "home_image.png")), size=(500, 250))

        self.txt_image= ct.CTkImage(light_image=Image.open(os.path.join(image_path, "txt_light.webp")), 
                                          dark_image=Image.open(os.path.join(image_path, "txt_dark.png")),size=(200, 200))

          # title
        self.navigation_frame_label = ct.CTkLabel(self.navigation_frame, text=" Ordonnancement", image=self.logo_image,
                                                             compound="left", font=ct.CTkFont(size=15, weight="bold"))
        self.navigation_frame_label.grid(row=0, column=0)#, padx=20, pady=20)
        
                #home button
        self.home_button = ct.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Home",
                                                   fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                                   image=self.home_image, anchor="w", command=self.home_button_event)
        self.home_button.grid(row=1, column=0)#, sticky="ew")
                # matrix button
        self.matrix_btn = ct.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Insertion des données",
                                                      fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                                      image=self.matrix_image, anchor="w", command=self.matrix_button_event)
        self.matrix_btn.grid(row=2, column=0)#, sticky="ew")
                # sequence button
        self.seq_btn = ct.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text=" Trouver la bonne sequence",
                                                      fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                                      image=self.sequence_image, anchor="w", command=self.sequence_button_event)
        self.seq_btn.grid(row=3, column=0)#, sticky="ew")
                # combinations button
        self.combinations_btn = ct.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Combinations",
                                                      fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                                      image=self.combinations_image, anchor="w", command=self.combinations_button_event)
        self.combinations_btn.grid(row=4, column=0)#, sticky="ew")
                # Gantt chart button
        self.Ganttt_btn = ct.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Diagramme de Gantt",
                                                      fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                                      image=self.Gantt_image, anchor="w",command=self.Gantt_button_event)
        self.Ganttt_btn.grid(row=5, column=0)#, sticky="ew")
                #temp de preparation
        self.checkbox_prep = ct.CTkCheckBox(master=self.navigation_frame, text= "utiluser les temps de preparation",command=self.checkbox_prep_fun)
        self.checkbox_prep.grid(row=6, column=0, padx=5, pady=5, sticky='ew')
                #blockage
        self.checkbox_block = ct.CTkCheckBox(master=self.navigation_frame, text= "utiluser la condition de bloquage",command=self.checkbox_block_fun)
        self.checkbox_block.grid(row=7, column=0, padx=5, pady=5, sticky='ew')
                #noidel
        self.checkbox_noidel = ct.CTkCheckBox(master=self.navigation_frame, text= "utiluser la condition de non arret",command=self.checkbox_noidel_fun)
        self.checkbox_noidel.grid(row=8, column=0, padx=5, pady=5, sticky='ew')
                #nowit
        self.checkbox_nowait = ct.CTkCheckBox(master=self.navigation_frame, text= "utiluser la condition de non attendre",command=self.checkbox_nowait_fun)
        self.checkbox_nowait.grid(row=9, column=0, padx=5, pady=5, sticky='ew')
                # tardiness
        self.checkbox_tardiness = ct.CTkCheckBox(master=self.navigation_frame, text= "utiluser les retards 'tardiness'",command=self.checkbox_tardiness_fun)
        self.checkbox_tardiness.grid(row=10, column=0, padx=5, pady=5, sticky='ew')
                # alerts
        self.alerts = ct.CTkLabel(self.navigation_frame,text="Alerts")
        self.alerts.grid(row=11, column=0)#, sticky="ew")
                #apparence mode menu
        self.appearance_mode_menu = ct.CTkOptionMenu(self.navigation_frame, values=["Dark", "Light" , "System"],
                                                                command=self.change_appearance_mode_event)
        self.appearance_mode_menu.grid(row=12, column=0, padx=5, pady=5, sticky='ew')
                #language menu
        #self.language_menu = ct.CTkOptionMenu(self.navigation_frame, values=["Français", "English"],
        #                                                        command=self.fun)
        #self.language_menu.grid(row=10, column=0, padx=20, pady=20, sticky='ew')

                    # create home frame
        self.home_frame = ct.CTkFrame(self, corner_radius=0, fg_color="transparent")
        self.home_frame.grid_columnconfigure(0, weight=1)

        self.home_frame_large_image_label = ct.CTkLabel(self.home_frame, text="", image=self.in_home_image)
        self.home_frame_large_image_label.grid(row=0, column=0, padx=40, pady=20)

        self.home_frame_large_label = ct.CTkLabel(self.home_frame, text="Créer par Salah Eddine Maimouni et Aya Kraimi\nDans cette application vous pouver faire\n la résolution des problèmes de production de type \nFlowshop où vous pouver ordonancer jobs sur des machines en série\n avec ou sans temps de préparation et avec ou sans blocage\navec condition de non attendre ou condition de non arret \n séparateur dicimale '.'\n séparateur dans les matrices et les séquences ','\n séparateur entre les matrices des préparation '-'\n 1) Entrer les donnée nécessaires \n 'temps de processus / temps de préparation / temps de delais / séquence'\n 2) vous pouver trouver la séquence optimisant Cmax on utilusant l'algorithme de jonhson ou CDS\n au cas d'avec aucune condition et sans temps de préparation\n 3) Vous pouver utiluser la fonction combinaisons soit pour\n tester des séquences spécifiques  ou pour trouver les séquences optimisant \n Cmax, Total Tardiness ou les deux \n4) la fonction Gantt permet de tracer le diagramme de Gantt \n et voir les performances de machines et d'autres detailes \n 5) Vous pouver extraire un rapport word daitailler \noù vous trouver toute étape de résolution de problème    ")
        self.home_frame_large_label.grid(row=1, column=0, padx=40, pady=20)


            # create matrix frame
        self.matrix_frame = ct.CTkFrame(self, corner_radius=0, fg_color="transparent")
        self.matrix_frame.grid_columnconfigure(0, weight=1)


        # ajouter la tabview
        self.tabview = ct.CTkTabview(self.matrix_frame)
        self.tabview.grid(row=0, column=0, columnspan=2)#, padx=40, pady=20)
        #enter text matrix
        self.tabview.add("selectionner les fichiers txt")
        self.text_insert = ct.CTkLabel(self.tabview.tab("selectionner les fichiers txt"),image = self.txt_image,text="")
        self.text_insert.pack()
        self.text_insert_button1 = ct.CTkButton(self.tabview.tab("selectionner les fichiers txt"), text="temps des processus",command=self.txt_insert_event)
        self.text_insert_button1.pack(side=ct.LEFT)
        self.text_insert_button2 = ct.CTkButton(self.tabview.tab("selectionner les fichiers txt"), text="temps de preparation",command=self.prep_insert_event)
        self.text_insert_button2.pack(side=ct.RIGHT)
        #enter matrix
        """self.tabview.add("inserer la matrice")
        self.matrix_insert = ct.CTkLabel(self.tabview.tab("inserer la matrice"),image = self.matrix_image,text="")
        self.matrix_insert.pack()
        self.not_fic = ct.CTkLabel(self.tabview.tab("inserer la matrice"),text="clique pour ajouter la matrice")
        self.not_fic.pack()
        self.matrix_insert_button = ct.CTkButton(self.tabview.tab("inserer la matrice"), text="click",command=self.matrix_insert_event)
        self.matrix_insert_button.pack()"""
        self.matrix_text = ct.CTkTextbox(self.matrix_frame, width= 200, height= 220)
        self.matrix_text.grid(row=2, column=0, padx=20, pady=20)
        self.prep_text = ct.CTkTextbox(self.matrix_frame, width= 200, height= 220)
        self.prep_text.grid(row=2, column=1, padx=20, pady=20)

        self.save_txt_button = ct.CTkButton(self.matrix_frame, text="enregistrer \n les temps de processus",command=self.save_txt_event)
        self.save_txt_button.grid(row=3, column=0)

        self.save_prep_button = ct.CTkButton(self.matrix_frame, text="enregistrer \n les temps de preparation",command=self.save_prep_event)
        self.save_prep_button.grid(row=3, column=1)
        self.dj_label = ct.CTkLabel(self.matrix_frame, text="temps de délais 'dj' :")
        self.dj_entry = ct.CTkEntry(self.matrix_frame)
        self.dj_label.grid(row=4, column=0,padx=5, pady=5)
        self.dj_entry.grid(row=4, column=1,padx=5, pady=5)

            # create combinations frame
        self.combinations_frame = ct.CTkFrame(self, corner_radius=0, fg_color="transparent")

        self.combination_test_btn= ct.CTkButton(self.combinations_frame, text="tester ces combinaisons",command=self.combinations_test_event)
        self.combination_test_btn.grid(row=0,column=0,columnspan=2)
        
        self.label_comb1 = ct.CTkLabel(self.combinations_frame, text="enter les combinaisons à tester")
        self.label_comb1.grid(row=1,column=0)
        self.combinations_entry_result = ct.CTkTextbox(self.combinations_frame, width= 200, height= 220)
        self.combinations_entry_result.grid(row=2,column=0, padx=20, pady=20)
        self.label_comb2 = ct.CTkLabel(self.combinations_frame, text="resultas de test")
        self.label_comb2.grid(row=1,column=1)
        self.combinations_test_result = ct.CTkTextbox(self.combinations_frame, width= 200, height= 220)
        self.combinations_test_result.grid(row=2,column=1, padx=20, pady=20)

        self.combination_btn= ct.CTkButton(self.combinations_frame, text="séquences avec \n minimisation de :",command=self.combinations_event)
        self.combination_btn.grid(row=3,column=0, padx=5, pady=10)
        self.critere_menu = ct.CTkOptionMenu(self.combinations_frame, values=["Cmax"])
        self.critere_menu.grid(row=3, column=1, padx=5, pady=10)
        self.combinations_result = ct.CTkTextbox(self.combinations_frame, width= 300, height= 220)
        self.combinations_result.grid(row=4,column=0,columnspan=2)

            # create sequence frame
        self.sequence_frame = ct.CTkFrame(self, corner_radius=0, fg_color="transparent")

        self.Jonhson_button = ct.CTkButton(self.sequence_frame, text="Jonhson",command=self.Jonhson_event)
        self.Jonhson_button.pack(padx=5, pady=3)
        self.jonhson_result = ct.CTkTextbox(self.sequence_frame)
        self.jonhson_result.pack(padx=5, pady=3)

        self.CDS_btn= ct.CTkButton(self.sequence_frame, text="CDS",command=self.CDS_event)
        self.CDS_btn.pack(padx=5, pady=3)
        self.CDS_label = ct.CTkLabel(self.sequence_frame, text="K =")
        self.CDS_label.pack(padx=5, pady=3)
        self.k_entry = ct.CTkEntry(self.sequence_frame)
        self.k_entry.pack(padx=5, pady=3)
        self.checkbox_k = ct.CTkCheckBox(master=self.sequence_frame, text= "tous les k")#, variable = self.long_var)
        self.checkbox_k.pack(padx=5, pady=3)
        self.CDS_result = ct.CTkTextbox(self.sequence_frame)
        self.CDS_result.pack(padx=5, pady=3)
            # create Gantt frame
        self.Gantt_frame = ct.CTkFrame(self, corner_radius=0, fg_color="transparent")
        self.seq_label = ct.CTkLabel(self.Gantt_frame, text="sequence = ")
        self.seq_label.pack(padx=5, pady=10)
        self.seq_entry = ct.CTkEntry(self.Gantt_frame)
        self.seq_entry.pack(padx=5, pady=10)
        self.Gantt_btn= ct.CTkButton(self.Gantt_frame, text="Diagramme de Gantt",command=self.Gantt_event)
        self.Gantt_btn.pack(padx=5, pady=10)
        self.report_frame = ct.CTkFrame(self.Gantt_frame, corner_radius=0)
        self.report_frame.pack(padx=120, pady=20,fill=ct.BOTH, expand=1)
        self.label_report = ct.CTkLabel(self.report_frame, text = "Génération de rapport:")
        self.label_report.pack()
        self.checkbox_regular_report = ct.CTkCheckBox(self.report_frame, text= "PFSP \n(Permutation Flow Shop Scheduling Problem)",command=self.checkbox_prep_fun)
        self.checkbox_regular_report.pack(padx=0, pady=3)   
        self.checkbox_prep_report = ct.CTkCheckBox(self.report_frame, text= "PFSP-SDST \n(Permutation Flow Shop Scheduling Problem \n with Sequence Depending Setup Time)",command=self.checkbox_prep_fun)
        self.checkbox_prep_report.pack(padx=0, pady=3)
        self.checkbox_block_report = ct.CTkCheckBox(master=self.report_frame, text= "BFSP \n(Blocking Flow Shop Scheduling Problem)",command=self.checkbox_block_fun)
        self.checkbox_block_report.pack(padx=0, pady=3)
        self.checkbox_prep_block_report = ct.CTkCheckBox(master=self.report_frame, text= "BFSP-SDST \n(Blocking Flow Shop Scheduling Problem \n with Sequence Depending Setup Time)",command=self.checkbox_block_fun)
        self.checkbox_prep_block_report.pack(padx=0, pady=3)
        self.checkbox_noidel_report = ct.CTkCheckBox(master=self.report_frame, text= "NIPFSP \n(No-Idel Permutation Flow Shop Scheduling Problem)",command=self.checkbox_noidel_fun)
        self.checkbox_noidel_report.pack(padx=0, pady=3)
        self.checkbox_nowait_report = ct.CTkCheckBox(master=self.report_frame, text= "NWPFSP \n(No-Wait Permutation Flow Shop Scheduling Problem)",command=self.checkbox_nowait_fun)
        self.checkbox_nowait_report.pack(padx=0, pady=3)
        self.checkbox_tardiness_report = ct.CTkCheckBox(master=self.report_frame, text= "Utiluser les retards 'tardiness'",command=self.checkbox_tardiness_fun)
        self.checkbox_tardiness_report.pack(padx=0, pady=3)
        self.checkbox_choosed_seq_report = ct.CTkCheckBox(master=self.report_frame, text= "Utiluser la séquence spécifique'",command=self.checkbox_tardiness_fun)
        self.checkbox_choosed_seq_report.pack(padx=0, pady=3)
        self.report_btn= ct.CTkButton(self.report_frame, text="générer un document",command=self.report_event)
        self.report_btn.pack(padx=5, pady=3)
        """
        self.report_frame.grid(row=0, column=0, sticky="nsew")
        self.report_frame.grid_rowconfigure(4, weight=1)"""
        #self.ctk_textbox_scrollbar = ct.CTkScrollbar(self.Gantt_frame)
        #self.ctk_textbox_scrollbar.pack(side=ct.LEFT)




            # select default frame
        self.select_frame_by_name("home")


    def fun(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")

    def change_appearance_mode_event(self, new_appearance_mode):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        ct.set_appearance_mode(new_appearance_mode)

    def select_frame_by_name(self, name):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        # show selected frame
        if name == "home":
            self.home_frame.grid(row=0, column=1, sticky="nsew")
        else:
            self.home_frame.grid_forget()
        if name == "matrix":
            self.matrix_frame.grid(row=0, column=1, sticky="nsew")
        else:
            self.matrix_frame.grid_forget()
        if name == "combinations":
            self.combinations_frame.grid(row=0, column=1, sticky="nsew")
        else:
            self.combinations_frame.grid_forget()
        if name == "sequence":
            self.sequence_frame.grid(row=0, column=1, sticky="nsew")
        else:
            self.sequence_frame.grid_forget()
        if name == "Gantt":
            self.Gantt_frame.grid(row=0, column=1, sticky="nsew")
        else:
            self.Gantt_frame.grid_forget()

    def home_button_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        self.select_frame_by_name("home")
        
    def matrix_button_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        self.select_frame_by_name("matrix")

    def combinations_button_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        self.select_frame_by_name("combinations")

    def sequence_button_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        self.select_frame_by_name("sequence")

    def Gantt_button_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        self.select_frame_by_name("Gantt")

    def txt_insert_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        path = get_file_path()
        print(path)
        f = open(path,"r")
        a=""
        for l in f:
            a=a+l+"\n"
        print(a)
        if type(ttm(a))==str:
            self.alerts.configure(text='le fichier est endomagé', fg_color="orange")
            return None
        self.matrix_text.delete(1.0,ct.END)
        self.matrix_text.insert(1.0,a)

    def prep_insert_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        path = get_file_path()
        print(path)
        f = open(path,"r")
        a=""
        for l in f:
            a=a+l+"\n"
        print(a)
        if type(char_to_lst(a))==str:
            self.alerts.configure(text='le fichier est endomagé', fg_color="orange")
            return None
        self.prep_text.delete(1.0,ct.END)
        self.prep_text.insert(1.0,a)

    """def matrix_insert_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        root = tk.Tk()
        window = MatrixEditor(root)
        root.mainloop()"""

    def save_txt_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        var = ttm(self.matrix_text.get(1.0,ct.END))
        if type(var)==str:
            self.alerts.configure(text=var, fg_color="orange")
            return None
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
        print(var)
        if file_path:
            with open(file_path, 'w') as file:
                file.write(self.matrix_text.get(1.0,ct.END))

    def save_prep_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        var = char_to_lst(self.prep_text.get(1.0,ct.END))
        if type(var)==str:
            self.alerts.configure(text='erreur de syntaxt', fg_color="orange")
            return None
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
        print(var)
        if file_path:
            with open(file_path, 'w') as file:
                file.write(self.prep_text.get(1.0,ct.END))

    def Jonhson_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        if self.checkbox_prep.get() == 1:
            self.alerts.configure(text='jonhson fonctionne seulement \n sans temps de preparation', fg_color="black")
            self.jonhson_result.delete(1.0,ct.END)
            return None
        if self.checkbox_block.get() == 1:
            self.alerts.configure(text='jonhson fonctionne seulement \n sans condition de blocage', fg_color="black")
            self.jonhson_result.delete(1.0,ct.END)
            return None
        if self.checkbox_noidel.get() == 1:
            self.alerts.configure(text='jonhson fonctionne seulement \n sans condition de non arret', fg_color="black")
            self.jonhson_result.delete(1.0,ct.END) 
            return None
        if self.checkbox_nowait.get() == 1:
            self.alerts.configure(text='jonhson fonctionne seulement \n sans condition de non attendre', fg_color="black")  
            self.jonhson_result.delete(1.0,ct.END)    
            return None
        matrix = ttm(self.matrix_text.get(1.0,ct.END))
        if type(matrix) == str:
            self.alerts.configure(text=matrix, fg_color="red")
            return None
        if matrix==[]:
            self.alerts.configure(text="la matrice n'existe pas", fg_color="red")
            return None 
        Ord = ordonnancement(M=matrix)
        out = Ord.Johnson()
        result = ""
        char_lst=["U = ","V = ","U_SPT = ","V_LPT = "]
        if len(matrix) != 2 and matrix != []:
            print(matrix)
            self.alerts.configure(text="Jonhson Algo ne fonctionne \n qu'avec deux machines ", fg_color="red")
            return None
        try:
            for o,char in zip(out,char_lst):
                result = result +char+str(o)
                result +="\n"
        except:
            self.alerts.configure(text="la matrice n'est pas exacte", fg_color="red")
            return None  
        L='\n'
        for lst in out[4]:
            char=""
            for i in range(len(lst)):
                if i!=len(lst)-1:
                    char+=str(lst[i])+","
                else:
                    char+=str(lst[i])
            L=L+char+'\n'
        result = result+"[UV] = "+L
        print(result)
        self.jonhson_result.delete(1.0,ct.END)
        self.jonhson_result.insert(1.0,result)

    def CDS_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        try:
            if self.checkbox_k.get()==0:
                int(self.k_entry.get())
        except ValueError:
            self.alerts.configure(text="k n'est pas valide", fg_color="red")
            return None
        if self.checkbox_prep.get() == 1:
            self.alerts.configure(text='CDS fonctionne seulement \n sans temps de preparation', fg_color="black")
            self.CDS_result.delete(1.0,ct.END)
            return None
        if self.checkbox_block.get() == 1:
            self.alerts.configure(text='CDS fonctionne seulement \n sans condition de blocage', fg_color="black")
            self.CDS_result.delete(1.0,ct.END)
            return None
        if self.checkbox_noidel.get() == 1:
            self.alerts.configure(text='CDS fonctionne seulement \n sans condition de non arret', fg_color="black")
            self.CDS_result.delete(1.0,ct.END) 
            return None
        if self.checkbox_nowait.get() == 1:
            self.alerts.configure(text='CDS fonctionne seulement \n sans condition de non attendre', fg_color="black")
            self.CDS_result.delete(1.0,ct.END)      
            return None
        matrix = ttm(self.matrix_text.get(1.0,ct.END))
        if type(matrix) == str:
            self.alerts.configure(text=matrix, fg_color="red")
            return None 
        if matrix==[]:
            self.alerts.configure(text="la matrice n'existe pas", fg_color="red")
            return None 
        if self.checkbox_k.get()==1:
            result = ""
            for k in range(1,len(matrix)):
                Ord = ordonnancement(M=matrix)
                out = Ord.CDS(k)
                L='\n'
                try:
                    for lst in out[4]:
                        char=""
                        for i in range(len(lst)):
                            if i!=len(lst)-1:
                                char+=str(lst[i])+","
                            else:
                                char+=str(lst[i])
                        L=L+char+'\n'
                except TypeError:
                    self.alerts.configure(text="CDS ne fonctionne qu'avec \n plus de deux machines", fg_color="red")
                    return None
                result = result + "k=" + str(k) + "->" + L +"\n"
            self.CDS_result.delete(1.0,ct.END)
            self.CDS_result.insert(1.0,result)
        else:
            if int(self.k_entry.get()) in range(1,len(matrix)):
                Ord = ordonnancement(M=matrix)
                try:
                    out = Ord.CDS(int(self.k_entry.get()))
                except:
                    self.alerts.configure(text="la matrice n'est pas exacte", fg_color="red")
                    return None 
                result = "pour k = "+str(self.k_entry.get())+"\n"
                char_lst=["U = ","V = ","U_SPT = ","V_LPT = "]
                try:
                    for o,char in zip(out,char_lst):
                        result = result +char+str(o)
                        result +="\n"
                except TypeError:
                    self.alerts.configure(text="CDS ne fonctionne qu'avec \n plus de deux machines", fg_color="red")
                    return None
                L='\n'
                for lst in out[4]:
                    char=""
                    for i in range(len(lst)):
                        if i!=len(lst)-1:
                            char+=str(lst[i])+","
                        else:
                            char+=str(lst[i])
                    L=L+char+'\n'
                result = result+"[UV] = "+L
                print(result)
                self.CDS_result.delete(1.0,ct.END)
                self.CDS_result.insert(1.0,result)
            else:
                self.alerts.configure(text="la valeur de k n'est pas exacte", fg_color="red")
                return None 

    def combinations_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        matrix = ttm(self.matrix_text.get(1.0,ct.END))
        matrix_prep = char_to_lst(self.prep_text.get(1.0,ct.END))
        if type(matrix) == str:
            self.alerts.configure(text=matrix, fg_color="red")
            return None
        if type(matrix_prep) == str:
            self.alerts.configure(text=matrix_prep, fg_color="red")
            return None
        if matrix_prep==[]:
            self.alerts.configure(text="les temps de preparation \n n'existe pas", fg_color="red")
            return None 
        if matrix==[]:
            self.alerts.configure(text="la matrice n'existe pas", fg_color="red")
            return None 
        if len(matrix[0])>9:
            self.alerts.configure(text="pour plus de 9 jobs cette fonction \n prend un tres long temps", fg_color="red")
            return None 
        DJ = dj_test(self.dj_entry.get())
        if self.checkbox_tardiness.get()==1:
            if type(DJ)==str:
                self.alerts.configure(text=DJ, fg_color="red")
                return None
            elif len(DJ)!=len(matrix[0]):
                self.alerts.configure(text="erreur dans les temps de delai", fg_color="red")
                return None
        result = ""
        if self.checkbox_prep.get() == 0 and self.checkbox_block.get() == 0 and self.checkbox_noidel.get() == 0 and self.checkbox_nowait.get() == 0:
            result+=" sans preparation \n et sans conditions \n"
        elif self.checkbox_prep.get() == 1 and self.checkbox_block.get() == 0 and self.checkbox_noidel.get() == 0 and self.checkbox_nowait.get() == 0:
            result+=" avec preparation \n et sans conditions \n"
        elif self.checkbox_prep.get() == 0 and self.checkbox_block.get() == 1 and self.checkbox_noidel.get() == 0 and self.checkbox_nowait.get() == 0:
            result+=" sans preparation \n et avec blocage \n"
        elif self.checkbox_prep.get() == 1 and self.checkbox_block.get() == 1 and self.checkbox_noidel.get() == 0 and self.checkbox_nowait.get() == 0:
            result+=" avec preparation \n et avec blocage \n"
        elif self.checkbox_prep.get() == 0 and self.checkbox_block.get() == 0 and self.checkbox_noidel.get() == 1 and self.checkbox_nowait.get() == 0:
            result+=" sans preparation \n et avec non arret \n"
        elif self.checkbox_prep.get() == 0 and self.checkbox_block.get() == 0 and self.checkbox_noidel.get() == 0 and self.checkbox_nowait.get() == 1:
            result+=" sans preparation \n et avec non attendre \n"
        Ord = ordonnancement(M=matrix,Mprep=matrix_prep,prep=self.checkbox_prep.get(),block=self.checkbox_block.get(),noidel=self.checkbox_noidel.get(),nowait=self.checkbox_nowait.get(),TT=self.checkbox_tardiness.get(),dj=DJ)
        if self.critere_menu.get()=="Cmax":
            out = Ord.choosed_combinations()
            result+="min(Cmaxs) = "+str(out[0])+"\n"
        elif self.critere_menu.get()=="TT":
            out = Ord.choosed_combinations_by_TT()
            result+="min(TTs) = "+str(out[0])+"\n"
        else:
            out = Ord.choosed_combinations_by_TT_Cmax()
            if type(out)==str:
                self.combinations_result.delete(1.0,ct.END)
                self.combinations_result.insert(1.0,out)
                return None
            else:
                result+="min(TTs) = "+str(out[0][0])+"\n min(Cmaxs) = "+str(out[0][1])+"\n"
        for seq in out[1]:
            char = ""
            for i in range(len(seq)):
                if i!=len(seq)-1:
                    char+=str(seq[i])+","
                else:
                    char+=str(seq[i])
            result = result + char
            result +="\n"
        self.combinations_result.delete(1.0,ct.END)
        self.combinations_result.insert(1.0,result)
    
    def combinations_test_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        matrix = ttm(self.matrix_text.get(1.0,ct.END))
        matrix_prep = char_to_lst(self.prep_text.get(1.0,ct.END))
        if type(matrix) == str:
            self.alerts.configure(text=matrix, fg_color="red")
            return None
        if type(matrix_prep) == str:
            self.alerts.configure(text=matrix_prep, fg_color="red")
            return None
        if matrix_prep==[]:
            self.alerts.configure(text="les temps de preparation \n n'existe pas", fg_color="red")
            return None 
        if matrix==[]:
            self.alerts.configure(text="la matrice n'existe pas", fg_color="red")
            return None
        DJ = dj_test(self.dj_entry.get())
        if self.checkbox_tardiness.get()==1:
            if type(DJ)==str:
                self.alerts.configure(text=DJ, fg_color="red")
                return None
            elif len(DJ)!=len(matrix[0]):
                self.alerts.configure(text="erreur dans les temps de delai", fg_color="red")
                return None
        seqqq= ttm(self.combinations_entry_result.get(1.0,ct.END))
        if type(seqqq) == str:
            self.alerts.configure(text=seqqq, fg_color="red")
            return None
        seq_to_test=[]
        for seq,k in zip(seqqq,range(len(seqqq))):
            seq_to_test.append([])
            for i in range(len(seqqq[0])):
                seq_to_test[k].append(int(seq[i]))
        if seq_to_test ==[]:
            self.alerts.configure(text="entrer les sequences à tester", fg_color="red")
            return None 
        if self.checkbox_prep.get() == 0 and self.checkbox_block.get() == 0 and self.checkbox_noidel.get() == 0 and self.checkbox_nowait.get() == 0:
            result=" sans preparation \n et sans conditions \n"
        elif self.checkbox_prep.get() == 1 and self.checkbox_block.get() == 0 and self.checkbox_noidel.get() == 0 and self.checkbox_nowait.get() == 0:
            result=" avec preparation \n et sans conditions \n"
        elif self.checkbox_prep.get() == 0 and self.checkbox_block.get() == 1 and self.checkbox_noidel.get() == 0 and self.checkbox_nowait.get() == 0:
            result=" sans preparation \n et avec blocage \n"
        elif self.checkbox_prep.get() == 1 and self.checkbox_block.get() == 1 and self.checkbox_noidel.get() == 0 and self.checkbox_nowait.get() == 0:
            result=" avec preparation \n et avec blocage \n"
        elif self.checkbox_prep.get() == 0 and self.checkbox_block.get() == 0 and self.checkbox_noidel.get() == 1 and self.checkbox_nowait.get() == 0:
            result=" sans preparation \n et avec non arret \n"
        elif self.checkbox_prep.get() == 0 and self.checkbox_block.get() == 0 and self.checkbox_noidel.get() == 0 and self.checkbox_nowait.get() == 1:
            result=" sans preparation \n et avec non attendre \n"
        Cmaxs=[]
        TTs=[]
        for seq in seq_to_test:
            Ord= ordonnancement(M=matrix,S=seq,Mprep=matrix_prep,prep=self.checkbox_prep.get(),block=self.checkbox_block.get(),noidel=self.checkbox_noidel.get(),nowait=self.checkbox_nowait.get(),TT=self.checkbox_tardiness.get(),dj=DJ)
            if self.checkbox_tardiness.get()==1:
                CT=Ord.Cmax_TT(seq)
                result+= str(seq)+"-> \n Cmax="+str(CT[0])+" / TT="+ str(CT[1])+"\n"
                Cmaxs.append(CT[0])
                TTs.append(CT[1])
            else:
                C=Ord.Cmax(seq)
                result+= str(seq)+"-> \n Cmax="+str(C)+"\n"
                Cmaxs.append(C)
        if self.checkbox_tardiness.get()==1:
            result= "min(Cmaxs)="+str(min(Cmaxs))+"\n"+"min(TTs)="+str(min(TTs))+"\n"+result
        else:
            result= "min(Cmaxs)="+str(min(Cmaxs))+"\n"+result
        self.combinations_test_result.delete(1.0,ct.END)
        self.combinations_test_result.insert(1.0,result)
        
    def Gantt_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        matrix = ttm(self.matrix_text.get(1.0,ct.END))
        matrix_prep = char_to_lst(self.prep_text.get(1.0,ct.END))
        if type(matrix) == str:
            self.alerts.configure(text=matrix, fg_color="red")
            return None
        if type(matrix_prep) == str:
            self.alerts.configure(text=matrix_prep, fg_color="red")
            return None
        if matrix_prep==[]:
            self.alerts.configure(text="les temps de preparation \n n'existe pas", fg_color="red")
            return None 
        if matrix==[]:
            self.alerts.configure(text="la matrice n'existe pas", fg_color="red")
            return None 
        DJ = dj_test(self.dj_entry.get())
        if self.checkbox_tardiness.get()==1:
            if type(DJ)==str:
                self.alerts.configure(text=DJ, fg_color="red")
                return None
            elif len(DJ)!=len(matrix[0]):
                self.alerts.configure(text="erreur dans les temps de delai", fg_color="red")
                return None
        seq = self.seq_entry.get()
        seq_lst = seq.split(",")
        job_lst = [] 
        for job in seq_lst:
            try:
                job_lst.append(int(job))
            except:
                self.alerts.configure(text="erreur dans la sequence", fg_color="red")
        if type(test_seq(job_lst,len(matrix[0])))==str:
            self.alerts.configure(text=test_seq(job_lst,len(matrix[0])), fg_color="red")
            return None
        Ord = ordonnancement(M=matrix,S=job_lst,Mprep=matrix_prep,prep=self.checkbox_prep.get(),
                             block=self.checkbox_block.get(),noidel=self.checkbox_noidel.get(),
                             nowait=self.checkbox_nowait.get(),dj=DJ,TT=self.checkbox_tardiness.get())
        TFR_TAR = Ord.TFR_TAR()
        TW = [[],[],[]]#Ord.TW()
        first_tabs = Ord.traite()
        """print(first_4_tabs[0])
        print(first_4_tabs[1])
        print(first_4_tabs[2])
        print(first_4_tabs[3])
        print(TFR_TAR)
        print(TW[0])
        print(TW[1])
        print(TW[2])"""



        task=Ord.Gantt()
        if self.checkbox_nowait.get() == 1:
            tables = Tableaux(first_tabs[0],first_tabs[1],[[]],[[]],TFR_TAR,first_tabs[2],[[]],[[]],[[]],task,job_lst,0,0,1,0,first_tabs[3],first_tabs[-1],first_tabs[-2],self.checkbox_tardiness.get())
        elif self.checkbox_noidel.get() == 1:
            tables = Tableaux(first_tabs[0],first_tabs[1],[[]],[[]],TFR_TAR,first_tabs[2],[[]],[[]],[[]],task,job_lst,0,0,0,1,first_tabs[3],first_tabs[-1],first_tabs[-2],self.checkbox_tardiness.get())
        elif self.checkbox_prep.get()==1 and self.checkbox_block.get()==0:
            TW = Ord.TW()
            tables = Tableaux(first_tabs[0],first_tabs[1],first_tabs[2],first_tabs[3],TFR_TAR,TW[0],TW[1],TW[2],[[]],task,job_lst,1,0,0,0,first_tabs[4],first_tabs[-1],first_tabs[-2],self.checkbox_tardiness.get())
        elif self.checkbox_prep.get()==0 and self.checkbox_block.get()==1:
            tables = Tableaux(first_tabs[0],first_tabs[1],[[]],[[]],TFR_TAR,first_tabs[4],[[]],[[]],[[]],task,job_lst,0,1,0,0,first_tabs[5],first_tabs[-1],first_tabs[-2],self.checkbox_tardiness.get())
        elif self.checkbox_prep.get()==0 and self.checkbox_block.get()==0:
            TW = Ord.TW()
            tables = Tableaux(first_tabs[0],first_tabs[1],[[]],[[]],TFR_TAR,TW[0],[[]],[[]],[[]],task,job_lst,0,0,0,0,first_tabs[2],first_tabs[-1],first_tabs[-2],self.checkbox_tardiness.get())
        elif self.checkbox_prep.get()==1 and self.checkbox_block.get()==1:
            tables = Tableaux(first_tabs[0],first_tabs[1],first_tabs[2],first_tabs[3],TFR_TAR,first_tabs[6],first_tabs[9],first_tabs[8],first_tabs[7],task,job_lst,1,1,0,0,first_tabs[10],first_tabs[-1],first_tabs[-2],self.checkbox_tardiness.get())
        if __name__ == "__main__":
            tables.mainloop()

        """result = "Cmax = " + str(Ord.Cmax(Seq=job_lst)) + "\n"
        for t1 , t2 , i in zip(TFR_TAR[0] , TFR_TAR[1] , range(len(TFR_TAR))):
            result += "TFR[" + str(i) + "] = " + str(round(t1*100,4)) + "%" + "\n"
            result += "TAR[" + str(i) + "] = " + str(round(t2*100,4)) + "%" + "\n"
        self.to_result.delete(1.0,ct.END)
        self.to_result.insert(1.0,result)"""
        #print(Ord.TW())

    def checkbox_prep_fun(self):
        if self.checkbox_prep.get()==1:
            self.checkbox_noidel.deselect()
            self.checkbox_nowait.deselect()
    def checkbox_block_fun(self):
        if self.checkbox_block.get()==1:
            self.checkbox_noidel.deselect()
            self.checkbox_nowait.deselect()
    def checkbox_noidel_fun(self):
        if self.checkbox_noidel.get()==1:
            self.checkbox_prep.deselect()
            self.checkbox_block.deselect()
            self.checkbox_nowait.deselect()
    def checkbox_nowait_fun(self):
        if self.checkbox_nowait.get()==1:
            self.checkbox_prep.deselect()
            self.checkbox_block.deselect()
            self.checkbox_noidel.deselect()
    def checkbox_tardiness_fun(self):
        if self.checkbox_tardiness.get()==1:
            self.critere_menu.destroy()
            self.critere_menu = ct.CTkOptionMenu(self.combinations_frame, values=["Cmax", "TT" , "Cmax & TT"])
            self.critere_menu.grid(row=3, column=1, padx=5, pady=10)
        else:
            self.critere_menu.destroy()
            self.critere_menu = ct.CTkOptionMenu(self.combinations_frame, values=["Cmax"])
            self.critere_menu.grid(row=3, column=1, padx=5, pady=10)
    
    def report_event(self):
        self.alerts.configure(text='Alerts', fg_color="transparent")
        matrix = ttm(self.matrix_text.get(1.0,ct.END))
        matrix_prep = char_to_lst(self.prep_text.get(1.0,ct.END))
        if type(matrix) == str:
            self.alerts.configure(text=matrix, fg_color="red")
            return None
        if type(matrix_prep) == str:
            self.alerts.configure(text=matrix_prep, fg_color="red")
            return None
        if matrix_prep==[]:
            self.alerts.configure(text="les temps de preparation \n n'existe pas", fg_color="red")
            return None 
        if matrix==[]:
            self.alerts.configure(text="la matrice n'existe pas", fg_color="red")
            return None 
        DJ = dj_test(self.dj_entry.get())
        if self.checkbox_tardiness_report.get()==1:
            if type(DJ)==str:
                self.alerts.configure(text=DJ, fg_color="red")
                return None
            elif len(DJ)!=len(matrix[0]):
                self.alerts.configure(text="erreur dans les temps de delai", fg_color="red")
                return None
        else:
            DJ=[]

        if self.checkbox_choosed_seq_report.get()==0:
            print(matrix,matrix_prep,DJ)
            doc = Report(Matrix=matrix, TP=matrix_prep,dj=DJ,regular=self.checkbox_regular_report.get(),prep=self.checkbox_prep_report.get(),block=self.checkbox_block_report.get(),prep_block=self.checkbox_prep_block_report.get(),noidel=self.checkbox_noidel_report.get(),nowait=self.checkbox_nowait_report.get())
        else:
            seq = self.seq_entry.get()
            seq_lst = seq.split(",")
            job_lst = [] 
            for job in seq_lst:
                try:
                    job_lst.append(int(job))
                except:
                    self.alerts.configure(text="erreur dans la sequence", fg_color="red")
            if type(test_seq(job_lst,len(matrix[0])))==str:
                self.alerts.configure(text=test_seq(job_lst,len(matrix[0])), fg_color="red")
                return None
            doc = Report(Seq=job_lst,Matrix=matrix, TP=matrix_prep,dj=DJ,regular=self.checkbox_regular_report.get(),prep=self.checkbox_prep_report.get(),block=self.checkbox_block_report.get(),prep_block=self.checkbox_prep_block_report.get(),noidel=self.checkbox_noidel_report.get(),nowait=self.checkbox_nowait_report.get())
    


if __name__ == "__main__":
    app = App()
    app.mainloop()


