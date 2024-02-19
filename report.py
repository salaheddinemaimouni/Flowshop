from docx import Document
from tkinter import Tk, filedialog
from flow_shop import ordonnancement
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from docx.shared import Inches
import matplotlib.colors as mcolors

class Report:
    def __init__(self, regular=0, prep=0, block=0, prep_block=0, noidel=0, nowait=0, TT=0, Seq=[], dj=[], Matrix=[[]], TP=[[[]]]):
        self.regular = regular
        self.prep = prep
        self.block = block
        self.prep_block = prep_block 
        self.noidel = noidel
        self.nowait = nowait
        self.TT = TT
        self.seq = Seq
        self.dj = dj
        if self.dj==[]:
            self.use_TT =0
        else:
            self.use_TT =1
        self.matrix = Matrix
        self.TP = TP
        self.Title1 = 1
        self.Title2 = 1
        self.Table_num = 1
        self.fig_num = 1
        self.m = len(self.matrix)
        self.n = len(self.matrix[0])
        self.doc = Document()
        self.data_add()
        if self.regular == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+'. PFSP (Permutation Flow Shop Scheduling Problem)').style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.regular_add()
            self.regular_add_without_seq()
        if self.prep == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+". PFSP-SDST (Permutation Flow Shop Scheduling Problem with Sequence Depending Setup Time)").style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.prep_add()
            self.seq_choose(1,0,0,0)
            self.prep_add()
        if self.block == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+". BFSP (Blocking Flow Shop Scheduling Problem)").style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.block_add()
            self.seq_choose(0,1,0,0)
            self.block_add()
        if self.prep_block == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+". BFSP-SDST (Blocking Flow Shop Scheduling Problem with Sequence Depending Setup Time)").style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.prep_block_add()
            self.seq_choose(1,1,0,0)
            self.prep_block_add()
        if self.noidel == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+". NIPFSP \n(No-Idel Permutation Flow Shop Scheduling Problem)").style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.noidel_add()
            self.seq_choose(0,0,0,1)
            self.noidel_add()
        if self.nowait == 1:
            self.seq = Seq
            self.Title2 = 1
            self.doc.add_heading("    "+str(self.Title1)+". NWPFSP \n(No-Wait Permutation Flow Shop Scheduling Problem)").style = 'Heading1'
            self.Title1+=1
            if self.seq!=[]:
                self.nowait_add()
            self.seq_choose(0,0,1,0)
            self.nowait_add()
        self.file_path = get_save_path()
        self.save_doc()

    def data_add(self):
        self.doc.add_heading(str(self.Title1)+'. Data de Problème:', level=1).style = 'Heading1'
        self.Title1+=1
        self.doc.add_heading("    "+str(self.Title2)+'. Tableau de temps de processing').style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("P[i][S[j]]: matrice des temps de processing")
        self.table_P = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_P.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Temps de processing').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_P.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(self.matrix[i-1][j-1])
                elif i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "P"
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(j-1)
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
        if self.TP!=[[[]]]:
            self.doc.add_heading("    "+str(self.Title2)+'. Tableau de temps de préparation').style = 'Heading2'
            self.Title2+=1
            self.doc.add_paragraph("TP[i][S[j1]][S[j2]]: Tenseur des temps de préparation")
            self.table_TP=[]
            for i in range(self.m):
                self.table_TP.append(self.doc.add_table(rows=self.n+1, cols=self.n+1)) 
                self.table_TP[i].style = 'Table Grid'
                caption = self.doc.add_paragraph()
                caption.add_run('Table ' + str(self.Table_num) +  'Temps de préparation sur la machine '+str(i)).bold = True
                caption.style = 'Caption'
                self.Table_num+=1
                for j1 in range(self.n+1):
                    for j2 in range(self.n+1):
                        cell = self.table_TP[i].cell(j1, j2)
                        # Ensure there is at least one paragraph and run in the cell
                        if not cell.paragraphs:
                            cell.add_paragraph()
                        if not cell.paragraphs[0].runs:
                            cell.paragraphs[0].add_run()
                        
                        if j1 != 0 and j2 != 0:
                            cell.paragraphs[0].runs[0].text = str(self.TP[i][j1-1][j2-1])
                        elif j1 == 0 and j2 == 0:
                            cell.paragraphs[0].runs[0].text = "TP"+str(i)
                        elif j1 == 0 and j2 != 0:
                            cell.paragraphs[0].runs[0].text = 'J'+str(j2-1)
                        elif j1 != 0 and j2 == 0:
                            cell.paragraphs[0].runs[0].text = 'J'+str(j1-1)
        if self.dj != []:
            self.doc.add_heading("    "+str(self.Title2)+'. Liste des délais de chaque Job').style = 'Heading2'
            self.Title2+=1
            self.doc.add_paragraph("d[S[j]]: List des délais de chaque job")
            self.table_dj = self.doc.add_table(rows=2, cols=self.n+1)
            self.table_dj.style = 'Table Grid'
            caption = self.doc.add_paragraph()
            caption.add_run('Table ' + str(self.Table_num) + ': Délais de chaque job').bold = True
            caption.style = 'Caption'
            self.Table_num+=1
            for i in range(2):
                for j in range(self.n+1):
                    cell = self.table_dj.cell(i, j)
                    # Ensure there is at least one paragraph and run in the cell
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    if not cell.paragraphs[0].runs:
                        cell.paragraphs[0].add_run()
                    
                    if i == 1 and j != 0:
                        cell.paragraphs[0].runs[0].text = str(self.dj[j-1])
                    elif i == 0 and j != 0:
                        cell.paragraphs[0].runs[0].text = 'J'+str(j-1)
                    elif i == 0 and j == 0:
                        cell.paragraphs[0].runs[0].text = ''
                    elif i == 1 and j == 0:
                        cell.paragraphs[0].runs[0].text = 'd'
        if self.seq != []:
            self.doc.add_heading("    "+str(self.Title2)+'. Séquence spécifique des jobs').style = 'Heading2'
            self.Title2+=1
            self.doc.add_paragraph("S[j]: Séquence spécifique des jobs")
            self.table_seq = self.doc.add_table(rows=1, cols=self.n+1)
            self.table_seq.style = 'Table Grid'
            caption = self.doc.add_paragraph()
            caption.add_run('Table ' + str(self.Table_num) + ': Séquence spécifique des jobs').bold = True
            caption.style = 'Caption'
            self.Table_num+=1
            for i in range(1):
                for j in range(self.n+1):
                    cell = self.table_seq.cell(i, j)
                    # Ensure there is at least one paragraph and run in the cell
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    if not cell.paragraphs[0].runs:
                        cell.paragraphs[0].add_run()
                    
                    if j != 0:
                        cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                    elif j == 0:
                        cell.paragraphs[0].runs[0].text = 'S'
    
    def regular_add_without_seq(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,M=self.matrix,Mprep=self.TP)
        if self.m<=2:
            self.doc.add_heading("    "+str(self.Title2)+". Séquence par algorithme de Jonhson").style = 'Heading2'
            self.Title2+=1
            Results = Problem.Johnson()
            self.seq = Results[-1]
        else:
            self.doc.add_heading("    "+str(self.Title2)+". Séquence par algorithme de CDS").style = 'Heading2'
            self.Title2+=1
            combinations = []
            Cmaxs = []
            Matrixs = []
            for k in range(1,self.m):
                print("k=",k)
                Results = Problem.CDS(k)
                print(Results,'mmmmmmmmmmmmmmmmmmmm')
                Problem = ordonnancement(TT=self.use_TT,dj=self.dj,S=Results[0][-1],do_report=1,M=self.matrix,Mprep=self.TP)
                combination = []
                Cmax=[]
                for seq in Results[0][-1]:
                    print("seq=",seq)
                    Cmax.append(Problem.Cmax(seq))
                    combination.append(seq)
                ind = Cmax.index(min(Cmax))
                Cmaxs.append(Cmax[ind])
                combinations.append(combination[ind])
                Matrixs.append(Results[-1])
            index = Cmaxs.index(min(Cmaxs))
            self.seq = combinations[index]
            print(self.seq,"dsfdsfds")
            matrix = Matrixs[index]
            k = index+1
            Results = Problem.CDS(k)
            Results = [Results[0][0],Results[0][1],Results[0][2],Results[0][3],Results[0][4]]
            self.doc.add_paragraph("après le teste de tous les k possibles de k = 1,...,"+str(self.m-1)+"on a trouvé que le meilleur k qui nous donne Cmax le plus optimale est k="+str(k))
            self.doc.add_paragraph("P[i][S[j]]: matrice des temps de processing trouvée par CDS")
            self.table_P_CDS = self.doc.add_table(rows=3, cols=self.n+1)
            self.table_P_CDS.style = 'Table Grid'
            caption = self.doc.add_paragraph()
            caption.add_run('Table ' + str(self.Table_num) + ': Temps de processing des deux machines fictives pou k optimale').bold = True
            caption.style = 'Caption'
            self.Table_num+=1
            for i in range(3):
                for j in range(self.n+1):
                    cell = self.table_P_CDS.cell(i, j)
                    # Ensure there is at least one paragraph and run in the cell
                    if not cell.paragraphs:
                        cell.add_paragraph()
                    if not cell.paragraphs[0].runs:
                        cell.paragraphs[0].add_run()
                    
                    if i != 0 and j != 0:
                        cell.paragraphs[0].runs[0].text = str(matrix[i-1][j-1])
                    elif i == 0 and j == 0:
                        cell.paragraphs[0].runs[0].text = "P_CDS"
                    elif i == 0 and j != 0:
                        cell.paragraphs[0].runs[0].text = 'J'+str(j-1)
                    elif i != 0 and j == 0:
                        cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
        
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,M=self.matrix,Mprep=self.TP)
        self.doc.add_paragraph("S[j]: Séquence trouvée par l'algotrithme de Jonhson")
        self.doc.add_paragraph("U[j]: Jobs tel que P[0][S[j]]<P[1][S[j]]:"+str(Results[0]))
        self.doc.add_paragraph("V[j]: Jobs tel que P[0][S[j]]>=P[1][S[j]]:"+str(Results[1]))
        self.doc.add_paragraph("U_SPT[j]: U[j] avec Jobs ordonancer par SPT(short processing Time), temps de processing croissant:"+str(Results[2]))
        self.doc.add_paragraph("V_LPT[j]: V[j] avec Jobs ordonancer par LPT(long processing Time), temps de processing décroissant:"+str(Results[3]))
        self.doc.add_paragraph("Alors la séquence trouvée par Jonhson est la suivante ")
        self.table_seq_jonhson = self.doc.add_table(rows=1, cols=self.n+1)
        self.table_seq_jonhson.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Séquence trouvée par l'algotrithme de Jonhson").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(1):
            for j in range(self.n+1):
                cell = self.table_seq_jonhson.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif j == 0:
                    cell.paragraphs[0].runs[0].text = 'S'
        self.regular_add()

    def seq_choose(self,p,b,now,noi):
        print(p,"ppppppppppppp")
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,M=self.matrix,Mprep=self.TP,prep=p,block=b,noidel=noi,nowait=now)
        is_TT = 0
        if self.dj!=[]:
            choosed = Problem.choosed_combinations_by_TT_Cmax()
            is_TT=1
            if type(choosed)==str:
                choosed = Problem.choosed_combinations()
                is_TT=0
        else:
            choosed = Problem.choosed_combinations()
        print(choosed,'cccccccccccccccccccccc')
        self.seq = choosed[1][0]
        if is_TT==0:
            self.doc.add_heading("    "+str(self.Title2)+". Séquence par Test de tous les permutations possibles avec min(Cmax)").style = 'Heading2'
        else:
            self.doc.add_heading("    "+str(self.Title2)+". Séquence par Test de tous les permutations possibles avec min(Cmax) et min(TT)").style = 'Heading2'
        self.Title2+=1
        self.table_seq_comb = self.doc.add_table(rows=1, cols=self.n+1)
        self.table_seq_comb.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Séquence trouvée par l'algotrithme de Test de toutes les séquences").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(1):
            for j in range(self.n+1):
                cell = self.table_seq_comb.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif j == 0:
                    cell.paragraphs[0].runs[0].text = 'S'

    def regular_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de PFSP (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("C[i][S[j]]: matrice des dates de fin de chaque job dans chaque machine \n")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrice des dates des fins et des débuts de chaque job sur chaque machine").style = 'Heading2'
        self.Title2+=1
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine PFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])
                                
        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine PFSP \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine PFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines PFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[0][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[1][i-1]*100,2))+"%"
        Figure = plot_Performances(Performances)
        self.doc.add_picture(Figure, width=Inches(5)) 

        
        self.doc.add_heading("    "+str(self.Title2)+". Temps d'attedre de chaque job entre chaque deux machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TW: (waiting Time) temps d'attedre de chaque job entre chaque deux machines")
        waitting_times = Problem.TW()
        self.table_TW = self.doc.add_table(rows=self.m, cols=self.n+1)
        self.table_TW.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Temps d'attendre de chaque job entre chaque deux machines PFSP").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m):
            for j in range(self.n+1):
                cell = self.table_TW.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "TW" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)+'/'+str(i)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(waitting_times[0][i-1][j-1])
        
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),0,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))
    
    def prep_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP,prep=1)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de PFSP-SDST (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("C[i][S[j]]: matrice des dates de fin de chaque job dans chaque machine \nM[i][S[j]][S[j]]: Tenseur des temps de préparation\n ")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrice des dates des fins et des débuts de chaque job sur chaque machine").style = 'Heading2'
        self.Title2+=1
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine PFSP-SDST').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])
                                
        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine PFSP-SDST \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine PFSP-SDST').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines (machines non arretées en préparation)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)\nTAP: (teaux d'arret de préparation)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=4)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines PFSP-SDST (machines non arretées en préparation)').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(4):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i == 0 and j == 3:
                    cell.paragraphs[0].runs[0].text = 'TAP'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[2][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[3][i-1]*100,2))+"%"
                elif i != 0 and j == 3:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[4][i-1]*100,2))+"%"
        Figure = plot_Performances([Performances[2],Performances[3],Performances[4]])
        self.doc.add_picture(Figure, width=Inches(5))
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines (machines arretées en préparation)").style = 'Heading2'
        self.Title2+=1
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines PFSP-SDST (machines arretées en préparation)').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[2][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[3][i-1]*100,2))+"%"
        Figure = plot_Performances([Performances[0],Performances[1]])
        self.doc.add_picture(Figure, width=Inches(5))
        self.doc.add_heading("    "+str(self.Title2)+". Temps d'attedre de chaque job entre chaque deux machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TW: (waiting Time) temps d'attedre de chaque job entre chaque deux machines")
        waitting_times = Problem.TW()
        self.table_TW = self.doc.add_table(rows=self.m, cols=self.n+1)
        self.table_TW.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Temps d'attendre de chaque job entre chaque deux machines PFSP-SDST").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m):
            for j in range(self.n+1):
                cell = self.table_TW.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "TW" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)+'/'+str(i)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(waitting_times[0][i-1][j-1]) 
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),0,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))

    def block_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP,block=1)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de BFSP (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("D[i][S[j]]: matrice des dates de fin de chaque job avec son blocage dans chaque machine, aussi début des jobs dans la machine précédente \npour la matrice D l'indexation des machines réel est de 1 à "+str(self.m)+"\n")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrices D,C,F").style = 'Heading2'
        self.Title2+=1
        self.table_D = self.doc.add_table(rows=self.m+2, cols=self.n+1)
        self.table_D.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job avec son blocage dans chaque machine BFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+2):
            for j in range(self.n+1):
                cell = self.table_D.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "D" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[-2][i-1][j-1])

        self.doc.add_paragraph("C[i][S[j]]:dates de fin de chaque job dans chaque machine BFSP\n")
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine BFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])

        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine BFSP \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine BFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])

        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines BFSP ').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[0][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[1][i-1]*100,2))+"%"
        Figure = plot_Performances(Performances)
        self.doc.add_picture(Figure, width=Inches(5)) 
        self.doc.add_heading("    "+str(self.Title2)+". Temps de blocage de chaque job dans chaque machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TB: (Blocking Time) temps de blocage de chaque job dans chaque machines")
        self.table_TB = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_TB.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Temps de blocage de chaque job dans chaque machines BFSP").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_TB.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "TB" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[4][i-1][j-1])      
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),0,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))   

    def prep_block_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP,block=1,prep=1)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de BFSP-SDST (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("D[i][S[j]]: matrice des dates de fin de chaque job avec son blocage dans chaque machine, aussi début des jobs dans la machine précédente \npour la matrice D l'indexation des machines réel est de 1 à "+str(self.m)+" et l'indexation des jobs réel de 1 à "+str(self.n)+"\n")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrices D,C,F").style = 'Heading2'
        self.Title2+=1
        self.table_D = self.doc.add_table(rows=self.m+2, cols=self.n+2)
        self.table_D.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job avec son blocage dans chaque machine BFSP-SDST').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+2):
            for j in range(self.n+2):
                cell = self.table_D.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "D" 
                elif i == 0 and j != 0 and j!=1:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-2]+1)
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'J0'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[-2][i-1][j-1])

        self.doc.add_paragraph("C[i][S[j]]:dates de fin de chaque job dans chaque machine BFSP-SDST\n")
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine BFSP-SDST').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])

        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine BFSP-SDST \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine BFSP-SDST').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])

        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines (machines non arretées en préparation)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)\nTAP: (teaux d'arret de préparation)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=4)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines BFSP-SDST (machines non arretées en préparation)').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(4):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i == 0 and j == 3:
                    cell.paragraphs[0].runs[0].text = 'TAP'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[2][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[3][i-1]*100,2))+"%"
                elif i != 0 and j == 3:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[4][i-1]*100,2))+"%"
        Figure = plot_Performances([Performances[2],Performances[3],Performances[4]])
        self.doc.add_picture(Figure, width=Inches(5))
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines (machines arretées en préparation)").style = 'Heading2'
        self.Title2+=1
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines BFSP-SDST (machines arretées en préparation)').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[2][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[3][i-1]*100,2))+"%"
        Figure = plot_Performances([Performances[0],Performances[1]])
        self.doc.add_picture(Figure, width=Inches(5))
        self.doc.add_heading("    "+str(self.Title2)+". Temps de blocage de chaque job dans chaque machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TB: (Blocking Time) temps de blocage de chaque job dans chaque machines")
        self.table_TB = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_TB.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Temps de blocage de chaque job dans chaque machines BFSP-SDST").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_TB.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "TB" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[6][i-1][j-1])    
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),0,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))

    def noidel_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP,noidel=1)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de NIPFSP (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("C[i][S[j]]: matrice des dates de fin de chaque job dans chaque machine \n")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrice des dates des fins et des débuts de chaque job sur chaque machine").style = 'Heading2'
        self.Title2+=1
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine NIPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])
                                
        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine NIPFSP \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine NIPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines NIPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[0][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[1][i-1]*100,2))+"%"
        Figure = plot_Performances(Performances)
        self.doc.add_picture(Figure, width=Inches(5)) 
        self.doc.add_heading("    "+str(self.Title2)+". Temps d'attedre de chaque job entre chaque deux machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TW: (waiting Time) temps d'attedre de chaque job entre chaque deux machines")
        self.table_TW = self.doc.add_table(rows=self.m, cols=self.n+1)
        self.table_TW.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ": Temps d'attendre de chaque job entre chaque deux machines NIPFSP").bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m):
            for j in range(self.n+1):
                cell = self.table_TW.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "TW" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)+'/'+str(i)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[2][i-1][j-1])
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),0,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))

    def nowait_add(self):
        Problem = ordonnancement(TT=self.use_TT,dj=self.dj,do_report=1,S=self.seq,M=self.matrix,Mprep=self.TP,nowait=1)
        Results = Problem.traite()
        self.doc.add_heading("    "+str(self.Title2)+". Etapes de résolution de NWPFSP (Cmax/TFT...)").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("C[i][S[j]]: matrice des dates de fin de chaque job dans chaque machine \n")
        self.doc.add_paragraph(Results[-1]+"\n")
        self.doc.add_heading("    "+str(self.Title2)+". Matrice des dates des fins et des débuts de chaque job sur chaque machine").style = 'Heading2'
        self.Title2+=1
        self.table_C = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_C.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de fin de chaque job dans chaque machine NWPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_C.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "C" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[0][i-1][j-1])
                                
        self.doc.add_paragraph("F[i][S[j]]:dates de début de chaque job dans chaque machine NWPFSP \npour trouver les dates de débuts de chaque job: F[i][S[j]] = C[i][S[j]] - P[i][S[j]] \n")
        self.table_F = self.doc.add_table(rows=self.m+1, cols=self.n+1)
        self.table_F.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': dates de début de chaque job dans chaque machine NWPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(self.n+1):
                cell = self.table_F.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = "F" 
                elif i == 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = 'J'+str(self.seq[j-1])
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j != 0:
                    cell.paragraphs[0].runs[0].text = str(Results[1][i-1][j-1])
        self.doc.add_heading("    "+str(self.Title2)+". Performances des machines").style = 'Heading2'
        self.Title2+=1
        self.doc.add_paragraph("TFR: (teaux de fonctionnement réel)\nTAR: (teaux d'arret réel)")
        Performances = Problem.TFR_TAR()
        self.table_Performances = self.doc.add_table(rows=self.m+1, cols=3)
        self.table_Performances.style = 'Table Grid'
        caption = self.doc.add_paragraph()
        caption.add_run('Table ' + str(self.Table_num) + ': Pérformances des machines NWPFSP').bold = True
        caption.style = 'Caption'
        self.Table_num+=1
        for i in range(self.m+1):
            for j in range(3):
                cell = self.table_Performances.cell(i, j)
                # Ensure there is at least one paragraph and run in the cell
                if not cell.paragraphs:
                    cell.add_paragraph()
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].add_run()
                
                if i == 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = ""
                elif i == 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = 'TFR'
                elif i == 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = 'TAR'
                elif i != 0 and j == 0:
                    cell.paragraphs[0].runs[0].text = 'M'+str(i-1)
                elif i != 0 and j == 1:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[0][i-1]*100,2))+"%"
                elif i != 0 and j == 2:
                    cell.paragraphs[0].runs[0].text = str(round(Performances[1][i-1]*100,2))+"%"
        Figure = plot_Performances(Performances)
        self.doc.add_picture(Figure, width=Inches(5)) 
        self.doc.add_heading("    "+str(self.Title2)+". Diagramme de Gantt, ").style = 'Heading2'
        self.Title2+=1
        Figure = gantt_plot(Problem.Gantt(),1,'PFSP Cmax ='+str(Results[0][-1][-1]))
        self.doc.add_picture(Figure, width=Inches(5))

    def save_doc(self):
        self.doc.save(self.file_path)
        print(f"Word document created at: {self.file_path}")

def get_save_path():
    root = Tk()
    root.withdraw()  # Hide the main window

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")],
        title="Save Word Document As"
    )

    root.destroy()  # Close the Tkinter window

    return file_path

def plot_Performances(TFR_TAR_plt):
    categories = np.arange(len(TFR_TAR_plt[0]))
    bar_width = 0.25
    names=["TFR","TAR","TAP"]
    fig, ax = plt.subplots()
    plt.figure(figsize=(20, 20))
    for i, t_list in enumerate(TFR_TAR_plt):
        bars = ax.bar(categories + i * bar_width, t_list, label=names[i], width=bar_width, alpha=0.7)
        for bar, prob in zip(bars, t_list):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5, "", ha='center', va='bottom')

    ax.set_xlabel('Machines')
    ax.set_ylabel('Performances(%)')
    ax.set_title("temps de fonctionnement réel(TFR) et Temps d'arret réel/préparation (TAR/TAP)")
    ax.set_xticks(categories)
    legend = ax.legend(loc='lower right')
    legend.set_draggable(True)
    image_stream = BytesIO()
    fig.savefig(image_stream, format='png')
    image_stream.seek(0)
    return image_stream

def gantt_plot(tasks,nowait,name):

    fig, ax = plt.subplots()
    # Configuration de l'axe Y
    ax.set_yticks(range(len(tasks)))
    ax.set_yticklabels(tasks.index)

    # Flattening the lists of lists
    all_start_times = [time for times_list in tasks['Start'] for times in times_list for time in times]
    all_end_times = [time for times_list in tasks['End'] for times in times_list for time in times]

    # Configuration de l'axe X
    ax.set_xticks(np.arange(min(all_start_times), max(all_end_times) + 1, step=1))
    ax.set_xlim(min(all_start_times), max(all_end_times) + 1)

    # Tracé des tâches
    for i, task in enumerate(tasks.itertuples()):
        for start_times, end_times, color, label in zip(task.Start, task.End, task.Color, task.Label):
            for start, end, text in zip(start_times, end_times, label):
                # Check if the color is a valid color name or hexadecimal code
                try:
                    mcolors.to_rgba(color)
                except ValueError:
                    color = 'blue'  # Default to blue if an invalid color is provided
                ax.barh(i, end - start, left=start, color=color, alpha=0.7)
                text_x = start + (end - start) / 2
                text_y = i
                ax.text(text_x, text_y, label, ha='center', va='center', color='white', fontweight='bold')

                # Add an arrow at the end of each job except for the first one
                if i > 0 and nowait==1:
                    arrow_start = end
                    arrow_end = end  # The arrow origin is at the bottom of the job
                    ax.arrow(arrow_start, text_y-0.4, 0, -0.1, head_width=0.2, head_length=0.1, fc=color, ec=color)

    # Configuration des étiquettes
    ax.set_xlabel('Axe de temps')
    ax.set_title('Diagramme de Gantt('+name+')')
    plt.grid()
    image_stream = BytesIO()
    fig.savefig(image_stream, format='png')
    image_stream.seek(0)
    return image_stream


"""from docx import Document
import matplotlib.pyplot as plt
from io import BytesIO
from PIL import Image
from docx.shared import Inches
from tkinter import Tk, filedialog
import os

def save_matplotlib_figure_to_word(figure):
    # Save the Matplotlib figure as a PNG image
    image_stream = BytesIO()
    figure.savefig(image_stream, format='png')
    image_stream.seek(0)

    # Open the Word document
    doc = Document()

    # Add a heading
    doc.add_heading('Matplotlib Figure', level=1)

    # Add a paragraph
    doc.add_paragraph('Here is a Matplotlib figure:')

    # Add the Matplotlib figure as an image
    doc.add_picture(image_stream, width=Inches(5))  # Adjust width as needed

    # Get the file path from a dialog
    file_path = get_save_path()

    # Save the Word document
    doc.save(file_path)

def get_save_path():
    root = Tk()
    root.withdraw()  # Hide the main window

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")],
        title="Save Word Document As"
    )

    root.destroy()  # Close the Tkinter window

    return file_path

# Example usage
fig, ax = plt.subplots()
ax.plot([1, 2, 3, 4], [5, 6, 7, 8])

# Call the function with the Matplotlib figure
save_matplotlib_figure_to_word(fig)"""

"""
list_prep2 = [[[3,2,3,2,3,2],
               [2,2,2,3,2,3],
               [4,2,2,3,2,4],
               [3,3,3,5,4,3],
               [3,2,3,2,4,3],
               [2,4,3,4,3,2]],
              [[4,3,3,2,3,4],
               [2,3,2,4,2,3],
               [2,2,2,5,3,2],
               [3,2,3,3,5,4],
               [2,4,3,4,5,2],
               [3,2,3,4,5,3]],
              [[5,4,3,3,2,3],
               [3,2,2,4,2,4],
               [2,2,3,2,3,3],
               [3,2,3,3,2,3],
               [3,2,3,3,4,4],
               [2,2,3,2,3,3]]]
matrice2 = [[7,9,5,9,8,4],
            [3,8,4,5,4,9],
            [5,5,8,7,6,7]]
S = [4,2,0,3,1,5]

R = Report(Matrix=matrice2, TP=list_prep2,dj=S,regular=1,prep=1,block=1,prep_block=1,noidel=1,nowait=1)


from docx import Document
from docx.shared import RGBColor, Pt, Inches
from tkinter import Tk, filedialog

def get_save_path():
    root = Tk()
    root.withdraw()  # Hide the main window

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")],
        title="Save Word Document As"
    )

    root.destroy()  # Close the Tkinter window

    return file_path

def create_word_document():
    # Get user input for file path and name using file dialog
    file_path = get_save_path()

    # Check if the user canceled the file dialog
    if not file_path:
        print("User canceled the operation.")
        return

    # Create a new Word document
    doc = Document()

    # Add a title (Heading 1)
    doc.add_heading('Title 1', level=1)

    # Add a second-level title (Heading 2)
    doc.add_heading('Title 2', level=2)

    # Add a paragraph
    doc.add_paragraph('This is a sample paragraph.')

    # Add a table with modified style
    table = doc.add_table(rows=3, cols=3)

    # Set table style
    table.style = 'Table Grid'

    # Modify cell style
    for row in table.rows:
        for cell in row.cells:
            # Ensure there is at least one paragraph and run
            if not cell.paragraphs:
                cell.add_paragraph()

            if not cell.paragraphs[0].runs:
                cell.paragraphs[0].add_run()

            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            cell.paragraphs[0].runs[0].font.size = Pt(12)  # Set font size
            cell.paragraphs[0].alignment = 1  # Center align content vertically
            cell.paragraphs[0].paragraph_format.left_indent = Inches(0.1)  # Set left indent

    # Modify row and column style
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            # Ensure there is at least one paragraph and run
            if not cell.paragraphs:
                cell.add_paragraph()

            if not cell.paragraphs[0].runs:
                cell.paragraphs[0].add_run()

            # Set blue color for the first row and column
            if i == 0 or j == 0:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255)
                cell.paragraphs[0].runs[0].font.bold = True  # Make the text bold

    # Add content to the cells
    for i in range(1, 3):
        for j in range(1, 3):
            table.cell(i, j).paragraphs[0].runs[0].text = f'Value {i}-{j}'

    # Save the document
    doc.save(file_path)
    print(f"Word document created at: {file_path}")

if __name__ == "__main__":
    create_word_document()

"""